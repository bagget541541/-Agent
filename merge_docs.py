#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多Word文档整合 + 持卡建议生成

模式B核心脚本（LLM 智能合并版）：
1. 读取多个Word文档
2. 提取 H1/H2 结构，转为标准 JSON items
3. LLM 智能合并（去重 + 整合 + 分类 + 摘要）
4. 生成持卡建议
5. 调用 word-merger/generate_report.py 输出最终 Word

新流程：
  多份 Word → read_docx_content() 提取 H1/H2 结构
    → 转为标准 JSON items 格式
    → LLM 智能合并（去重 + 整合 + 综合分析）
    → 输出标准 JSON (CreditCardBatch)
    → word-merger/generate_report.py → 最终 Word
"""

import json
import os
import re
import sys
import hashlib
import subprocess
import argparse
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "word-merger" / "scripts"))
try:
    from utils import IMAGE_WIDTH_CM
except Exception:
    IMAGE_WIDTH_CM = 13

from common.config import KNOWN_FIELD_NAMES


# ═══════════════════════════════════════════════════════════════
#  Word 解析（保留原有三种解析模式）
# ═══════════════════════════════════════════════════════════════


def _detect_parse_mode(doc) -> str:
    """
    自动检测文档解析模式

    Returns:
        "standard" — Heading 1/2 层级清晰，原逻辑不变
        "announcement" — 只有 Heading 2 无 H1，公告模式
        "field" — 无 Heading，字段加粗占比 > 20%，字段模式
    """
    h1_count = 0
    h2_count = 0
    bold_field_count = 0
    total_para_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        total_para_count += 1
        style = para.style.name

        if style == 'Heading 1':
            h1_count += 1
        elif style == 'Heading 2':
            h2_count += 1

        # 检测段落是否以加粗字段名开头
        if para.runs:
            first_run = para.runs[0]
            if first_run.bold and first_run.text.strip():
                for field in KNOWN_FIELD_NAMES:
                    if text.startswith(field + '：') or text.startswith(field + ':'):
                        bold_field_count += 1
                        break

    # 模式切换逻辑
    # 有 H1 → 始终 standard（H1 结构优先于字段检测）
    if h1_count > 0:
        return "standard"

    if h2_count > 0:
        return "announcement"

    if total_para_count > 0 and bold_field_count / total_para_count > 0.2:
        return "field"

    return "standard"


def _parse_announcement_mode(doc, docx_path, images_map, ns) -> dict:
    """公告模式：无 H1，只有 H2 → 自动创建 H1 容器"""
    content = {
        "file": docx_path,
        "title": "",
        "h1_sections": [],
        "full_text": ""
    }

    # 自动创建"银行公告"作为 H1
    current_h1 = {
        "title": "银行公告",
        "h2_items": [],
        "content": []
    }
    content["h1_sections"].append(current_h1)
    current_h2 = None
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        # 检查段落中的图片
        drawings = para._element.findall('.//w:drawing', ns)
        para_images = []
        for d in drawings:
            blips = d.findall('.//a:blip', ns)
            for b in blips:
                rid = b.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rid and rid in images_map:
                    para_images.append(rid)

        if not text and not para_images:
            continue

        if style == 'Heading 2':
            current_h2 = {
                "title": text,
                "content": [],
                "images": []
            }
            current_h1["h2_items"].append(current_h2)
        elif current_h2 is not None:
            if text:
                # 尝试从文本中识别银行名，替换 H1 标题
                bank_match = re.search(
                    r'(农行|工行|建行|招行|中行|交行|浦发|中信|光大|民生|华夏|'
                    r'兴业|广发|平安|邮储|北京银行|上海银行|南京银行|宁波银行|'
                    r'江苏银行|杭州银行|秦农银行|农商银行|徽商银行|浙商银行|'
                    r'渤海银行|恒丰银行|银行)',
                    text
                )
                if bank_match and current_h1["title"] == "银行公告":
                    bank_name = bank_match.group(1)
                    if bank_name != "银行":
                        current_h1["title"] = bank_name + "公告"
                current_h2["content"].append(text)
            current_h2["images"].extend(para_images)
        elif current_h1 is not None:
            if text:
                current_h1["content"].append(text)

        if text:
            full_text_parts.append(text)

    content["full_text"] = "\n".join(full_text_parts)
    content["images_map"] = images_map
    return content


def _parse_field_mode(doc, docx_path, images_map, ns) -> dict:
    """字段模式：解析 KNOWN_FIELD_NAMES 结构化字段"""
    content = {
        "file": docx_path,
        "title": "",
        "h1_sections": [],
        "full_text": ""
    }

    current_h1 = None
    current_h2 = None
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()

        # 检查图片
        drawings = para._element.findall('.//w:drawing', ns)
        para_images = []
        for d in drawings:
            blips = d.findall('.//a:blip', ns)
            for b in blips:
                rid = b.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rid and rid in images_map:
                    para_images.append(rid)

        if not text and not para_images:
            continue

        # 检测字段名（优先用加粗检测，其次全文匹配）
        field_match = None
        first_run_bold = False

        if para.runs:
            first_run = para.runs[0]
            if first_run.bold and first_run.text.strip():
                first_run_bold = True
                first_text = first_run.text.strip()
                for field in KNOWN_FIELD_NAMES:
                    if first_text.startswith(field + '：') or first_text.startswith(field + ':'):
                        field_match = field
                        break

        # 全文兜底匹配
        if not field_match:
            for field in KNOWN_FIELD_NAMES:
                if text.startswith(field + '：') or text.startswith(field + ':'):
                    field_match = field
                    break

        if field_match:
            colon = '：' if '：' in text else ':'
            value = text.split(colon, 1)[1].strip()

            if field_match == '银行' or field_match == '发卡银行':
                # 创建 H1 容器，标题 = 银行名
                current_h1 = {
                    "title": value if value else "银行信息",
                    "h2_items": [],
                    "content": []
                }
                content["h1_sections"].append(current_h1)
                current_h2 = None
            elif field_match == '标题':
                # 创建 H2 条目
                current_h2 = {
                    "title": value if value else text,
                    "content": [],
                    "images": [],
                    "url": ""
                }
                if current_h1:
                    current_h1["h2_items"].append(current_h2)
                else:
                    # 无 H1 时创建默认容器
                    current_h1 = {
                        "title": "微信文章",
                        "h2_items": [],
                        "content": []
                    }
                    content["h1_sections"].append(current_h1)
                    current_h1["h2_items"].append(current_h2)
            elif field_match == '原文链接':
                if current_h2:
                    current_h2["url"] = value
                elif current_h1:
                    current_h1["content"].append(text)
            elif field_match in ('发布时间', '作者'):
                if current_h2:
                    if "metadata" not in current_h2:
                        current_h2["metadata"] = {}
                    current_h2["metadata"][field_match] = value
                elif current_h1:
                    current_h1["content"].append(text)
            else:
                # 其他字段：关键信息、点评、消息内容等 → 归入当前 H2 的 content
                if current_h2:
                    current_h2["content"].append(text)
                elif current_h1:
                    current_h1["content"].append(text)
        else:
            # 普通段落
            if current_h2:
                if text:
                    current_h2["content"].append(text)
                current_h2["images"].extend(para_images)
            elif current_h1:
                if text:
                    current_h1["content"].append(text)

        if text:
            full_text_parts.append(text)

    content["full_text"] = "\n".join(full_text_parts)
    content["images_map"] = images_map
    return content


def read_docx_content(docx_path: str) -> dict:
    """
    读取 Word 文档内容，自适应三种解析模式

    - standard: Heading 1/2 层级清晰（原逻辑不变）
    - announcement: 只有 H2 无 H1 → 自动创建 H1 容器
    - field: KNOWN_FIELD_NAMES 字段加粗 → 结构化解析

    Returns:
        包含层级章节的字典
    """
    try:
        from docx import Document

        doc = Document(docx_path)
        images_map = _extract_docx_images(docx_path)

        # XML 命名空间
        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }

        # 自动检测解析模式
        mode = _detect_parse_mode(doc)
        print(f"  [Parse Mode] {docx_path} → {mode}")

        if mode == "announcement":
            content = _parse_announcement_mode(doc, docx_path, images_map, ns)
        elif mode == "field":
            content = _parse_field_mode(doc, docx_path, images_map, ns)
        else:
            content = _parse_standard_mode(doc, docx_path, images_map, ns)

        return content

    except Exception as e:
        print(f"  [Error] Read failed {docx_path}: {e}")
        return {"file": docx_path, "title": Path(docx_path).stem, "h1_sections": [], "full_text": ""}


def _parse_standard_mode(doc, docx_path, images_map, ns) -> dict:
    """标准模式：原 Heading 1 → Heading 2 层级逻辑不变"""
    content = {
        "file": docx_path,
        "title": "",
        "h1_sections": [],
        "full_text": ""
    }

    full_text_parts = []
    current_h1 = None
    current_h2 = None

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        # 检查段落中的图片
        drawings = para._element.findall('.//w:drawing', ns)
        para_images = []
        for d in drawings:
            blips = d.findall('.//a:blip', ns)
            for b in blips:
                rid = b.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rid and rid in images_map:
                    para_images.append(rid)

        # 跳过空段落（但保留有图片的段落）
        if not text and not para_images:
            continue

        # 识别 Heading 1
        if style == 'Heading 1':
            current_h1 = {
                "title": text,
                "h2_items": [],
                "content": []
            }
            content["h1_sections"].append(current_h1)
            current_h2 = None
        # 识别 Heading 2
        elif style == 'Heading 2' and current_h1 is not None:
            current_h2 = {
                "title": text,
                "content": [],
                "images": []
            }
            current_h1["h2_items"].append(current_h2)
        # 普通段落或列表
        elif current_h2 is not None:
            if text:
                current_h2["content"].append(text)
            current_h2["images"].extend(para_images)
        elif current_h1 is not None:
            if text:
                current_h1["content"].append(text)
        else:
            if not content["title"] and len(text) < 80:
                content["title"] = text

        if text:
            full_text_parts.append(text)

    content["full_text"] = "\n".join(full_text_parts)
    content["images_map"] = images_map
    return content


def _extract_docx_images(docx_path: str) -> dict:
    """提取 docx 中的所有图片，按 byte 哈希去重，返回 {rId: bytes} 映射"""
    import zipfile
    from lxml import etree

    images = {}
    seen_hashes = set()
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            # 读取 rels 文件获取 rId → media 映射
            rels_xml = zf.read('word/_rels/document.xml.rels')
            rels_root = etree.fromstring(rels_xml)
            ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}

            rid_to_target = {}
            for rel in rels_root.findall('r:Relationship', ns):
                rid = rel.get('Id')
                target = rel.get('Target')
                rel_type = rel.get('Type', '')
                if 'image' in rel_type and target:
                    rid_to_target[rid] = target

            # 提取图片文件，按 byte 哈希去重
            for rid, target in rid_to_target.items():
                img_path = f'word/{target}' if not target.startswith('/') else target[1:]
                try:
                    img_data = zf.read(img_path)
                    img_hash = hashlib.md5(img_data).hexdigest()
                    if img_hash not in seen_hashes:
                        seen_hashes.add(img_hash)
                        images[rid] = img_data
                    else:
                        # 重复图片，保留第一个出现的 rid 映射
                        for existing_rid, existing_data in images.items():
                            if hashlib.md5(existing_data).hexdigest() == img_hash:
                                # 建立别名映射 rId → 已有图片
                                images[rid] = existing_data
                                break
                except KeyError:
                    pass

    except Exception:
        pass

    return images


def _title_similarity(t1: str, t2: str) -> float:
    """
    计算两个标题的相似度（0.0 ~ 1.0），使用 Jaccard 字符二元组相似度

    阈值参考：同一银行同一条公告的标题相似度通常 ≥ 0.85
    """
    if not t1 or not t2:
        return 0.0
    # 提取非空格字符的二元组集合
    s1 = set(t1[i:i+2] for i in range(len(t1) - 1) if t1[i] != ' ')
    s2 = set(t2[i:i+2] for i in range(len(t2) - 1) if t2[i] != ' ')
    if not s1 or not s2:
        return 0.0
    intersection = s1 & s2
    union = s1 | s2
    return len(intersection) / len(union)


# ═══════════════════════════════════════════════════════════════
#  contents_to_items — H1/H2 结构 → 标准 JSON items
# ═══════════════════════════════════════════════════════════════


# H1 标题 → 标准分类映射
_H1_CATEGORY_MAP = {
    "新卡": "新卡", "新卡资讯": "新卡", "新卡发布": "新卡", "新卡推荐": "新卡",
    "权益变更": "权益变更", "权益调整": "权益变更", "权益变化": "权益变更",
    "活动": "活动", "优惠活动": "活动", "银行活动": "活动", "热门活动": "活动",
    "公告": "公告", "银行公告": "公告", "通知": "公告", "重要公告": "公告",
    "其他": "其他",
}


def categorize_h1(h1_title: str) -> str:
    """将 H1 标题映射到标准分类"""
    for key, cat in _H1_CATEGORY_MAP.items():
        if key in h1_title:
            return cat
    return "其他"


# 银行名正则
_BANK_RE = re.compile(
    r'(农业银行|工商银行|建设银行|招商银行|中国银行|交通银行|'
    r'浦发银行|中信银行|光大银行|民生银行|华夏银行|兴业银行|广发银行|'
    r'平安银行|邮储银行|北京银行|上海银行|南京银行|宁波银行|'
    r'江苏银行|杭州银行|秦农银行|农商银行|徽商银行|浙商银行|'
    r'渤海银行|恒丰银行|'
    r'农行|工行|建行|招行|中行|交行|'
    r'浦发|中信|光大|民生|华夏|兴业|广发|平安|邮储)'
)


def extract_bank_from_content(h1_title: str, content_lines: list) -> str:
    """从 H1 标题或 content 中提取银行名"""
    # 优先从 H1 标题提取
    m = _BANK_RE.search(h1_title)
    if m:
        return m.group(1)

    # 从 content 文本中提取
    text = " ".join(content_lines[:5])
    m = _BANK_RE.search(text)
    if m:
        return m.group(1)

    return ""


def extract_structured_fields(content_lines: list) -> dict:
    """从 content 文本中提取 '字段名：值' 格式的结构化字段"""
    structured = {}
    # 别名映射：将上游常见字段名映射为下游约定的标准键名
    alias_map = {
        '亮点': '卡亮点', '卡亮点': '卡亮点',
        '生效日期': '时间', '时间': '时间', '发布时间': '时间', '活动时间': '时间',
        '原文链接': '原文链接', '来源': '来源', '来源链接': '来源',
        '详情': '详情', '卡种': '卡种', '适用人群': '适用人群',
        '影响范围': '影响范围', '变更内容': '变更内容', '点评': '点评',
        '消息内容': '消息', '消息': '消息', '结论': '结论', '亮点摘要': '卡亮点'
    }

    for line in content_lines:
        line = line.strip()
        if not line:
            continue
        m = re.match(r'^(.{1,16})[：:]\s*(.+)', line)
        if m:
            field_name = m.group(1).strip()
            field_value = m.group(2).strip()
            if len(field_value) <= 2000:
                # 标准化字段名
                std_name = alias_map.get(field_name, field_name)
                # 如果字段已存在且已有值，则跳过以避免覆盖更完整的信息
                if std_name in structured and structured[std_name]:
                    continue
                structured[std_name] = field_value
    return structured


def resolve_image_paths(h2: dict, content: dict) -> list:
    """将 H2 条目的图片 RID 转为文件路径，写入磁盘供 generate_report.py 使用"""
    images_map = content.get("images_map", {})
    img_rids = h2.get("images", [])
    if not img_rids or not images_map:
        return []

    # 确保临时图片目录存在
    img_dir = PROJECT_ROOT / "data" / "images" / "merge_tmp"
    img_dir.mkdir(parents=True, exist_ok=True)

    paths = []
    for rid in img_rids:
        img_data = images_map.get(rid)
        if not img_data:
            continue

        # 确定扩展名
        ext = ".png"
        if img_data[:3] == b'\xff\xd8\xff':
            ext = ".jpg"
        elif img_data[:4] == b'\x89PNG':
            ext = ".png"
        elif img_data[:4] == b'GIF8':
            ext = ".gif"
        elif img_data[:4] == b'RIFF':
            ext = ".webp"

        # 用哈希命名，避免重复写入
        img_hash = hashlib.md5(img_data).hexdigest()[:12]
        img_file = img_dir / f"{rid}_{img_hash}{ext}"

        if not img_file.exists():
            img_file.write_bytes(img_data)

        paths.append(str(img_file))

    return paths


def contents_to_items(contents: list) -> list:
    """将 read_docx_content 输出的 H1/H2 结构转为标准 JSON items 格式"""
    items = []
    for content in contents:
        h1_sections = content.get("h1_sections", [])

        for h1 in h1_sections:
            h1_title = h1.get("title", "")
            category = categorize_h1(h1_title)

            for h2 in h1.get("h2_items", []):
                h2_title = h2.get("title", "").strip()
                content_lines = h2.get("content", [])

                # 跳过噪音条目
                raw_text = "\n".join(content_lines)
                if len(raw_text.strip()) < 5:
                    continue
                if "以上内容为广告" in raw_text:
                    continue

                bank = extract_bank_from_content(h1_title, content_lines)
                url = h2.get("url", "")

                structured = extract_structured_fields(content_lines)
                # 如果 item 顶层没有 url，尝试从 structured 的标准字段提升
                if not url:
                    for cand in ('原文链接', '来源'):
                        if structured.get(cand):
                            url = structured.get(cand)
                            break

                item = {
                    "category": category,
                    "title": h2_title,
                    "bank": bank,
                    "url": url,
                    "raw_text": raw_text[:3000],
                    "structured": structured,
                    "highlight_summary": "",
                    "images": resolve_image_paths(h2, content),
                    "_source_file": Path(content.get("file", "")).name,
                }
                items.append(item)

    return items


# ═══════════════════════════════════════════════════════════════
#  llm_merge — LLM 智能合并
# ═══════════════════════════════════════════════════════════════


def llm_merge(items: list, sources: list) -> dict:
    """
    LLM 智能合并：去重 + 整合 + 输出标准 JSON batch

    Args:
        items: contents_to_items 输出的标准 items 列表
        sources: 来源文件名列表

    Returns:
        {"batch_label": "...", "items": [...], "stats": {...}}
    """
    if not items:
        return {
            "batch_label": "信用卡资讯（合并）",
            "items": [],
            "stats": {"total_in": 0, "total_out": 0},
        }

    # 构造 items 摘要供 LLM 消费（精简：只发 title+bank+category+source）
    items_summary = []
    for i, item in enumerate(items):
        items_summary.append(f"  [{i}] {item.get('category', '')} | {item.get('bank', '')} | {item.get('title', '')} | {item.get('_source_file', '')}")

    system_msg = """你是信用卡资讯合并专家。给定多份周报的资讯条目索引列表，执行去重。

规则：
- 同一银行+同一标题/高度相似标题 → 保留一条（去重）
- 不同银行的不同活动/公告 → 保留
- 为每条生成一句话 highlight_summary

输出严格 JSON：
```json
{
  "items": [
    {
      "category": "新卡/权益变更/活动/公告/其他",
      "title": "条目标题",
      "bank": "银行名",
      "url": "原文链接（如有）",
      "structured": {},
      "highlight_summary": "一句话亮点摘要",
      "original_indices": [0, 10]
    }
  ],
  "stats": {
    "total_in": 原始条目数,
    "total_out": 合并后条目数,
    "merged": ["被合并的条目说明"],
    "dropped": ["被丢弃的条目说明"]
  }
}

注意：
- 保留所有不同银行的不同活动/公告，不要过度合并
- 只合并真正重复的条目（同一事项的不同来源报道）
- 如果两篇报道了同一事项但内容互补，选择信息更完整的一条
- 每个 item 的 structured 字段可以为空对象 {}
- original_indices: 必填，列出该条目对应的原始索引号（未合并的也必须填，如 [3]）

给定的条目索引映射：
""" + "\n".join(items_summary)

    user_msg = "请合并以上资讯条目，输出严格 JSON。原始条目数: " + str(len(items))

    # 尝试调用 LLM
    from common.llm_client import call_llm_simple
    content, error = call_llm_simple(
        system_msg, user_msg,
        temperature=0.3, max_tokens=16384, timeout=300,
    )

    if error or not content:
        print(f"  [Warn] LLM merge failed ({error})，使用基础去重")
        return _fallback_merge(items, sources)

    # 解析 LLM 输出 JSON
    try:
        # 尝试提取 JSON 代码块
        m = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", content, re.DOTALL)
        json_text = m.group(1).strip() if m else content.strip()

        # 如果 code block 提取失败，尝试直接提取 JSON 对象
        if not json_text or not json_text.startswith("{"):
            m2 = re.search(r"\{[\s\S]*\}", content)
            if m2:
                json_text = m2.group()

        result = json.loads(json_text)

        # 提取 items 和 stats
        llm_items = result.get("items", [])
        stats = result.get("stats", {})

        # 将 LLM 输出的 items 映射回完整结构
        final_items = []
        for li in llm_items:
            # 通过 original_indices 找到所有原始条目
            orig_indices = li.get("original_indices", [])
            originals = []
            if orig_indices:
                for idx in orig_indices:
                    if 0 <= idx < len(items):
                        originals.append(items[idx])
            # fallback: 用标题匹配找
            if not originals:
                orig = _find_original_item(li, items)
                if orig:
                    originals = [orig]

            # 合并所有原始条目的 raw_text / structured / images
            merged_raw = "\n".join(o.get("raw_text", "") for o in originals if o.get("raw_text"))
            merged_structured = {}
            for o in originals:
                for k, v in (o.get("structured") or {}).items():
                    if v and k not in merged_structured:
                        merged_structured[k] = v
            merged_images = []
            for o in originals:
                merged_images.extend(o.get("images", []))
            source_file = originals[0].get("_source_file", "") if originals else ""

            item = {
                "category": li.get("category", "其他"),
                "title": li.get("title", ""),
                "bank": li.get("bank", ""),
                "url": li.get("url", "") or (originals[0].get("url", "") if originals else ""),
                "raw_text": merged_raw or (originals[0].get("raw_text", "") if originals else ""),
                "structured": li.get("structured", {}) or merged_structured,
                "highlight_summary": li.get("highlight_summary", ""),
                "images": merged_images,
                "_source_file": source_file,
            }
            final_items.append(item)

        batch_label = f"信用卡资讯（{len(sources)}份合并）"
        print(f"  [LLM] {len(items)} → {len(final_items)} items (merged: {stats.get('merged', [])})")

        return {
            "batch_label": batch_label,
            "items": final_items,
            "stats": {
                "total_in": len(items),
                "total_out": len(final_items),
                "sources": sources,
                **stats,
            },
        }

    except (json.JSONDecodeError, KeyError) as e:
        print(f"  [Warn] LLM merge JSON parse failed: {e}，使用基础去重")
        return _fallback_merge(items, sources)


def _find_original_item(llm_item: dict, original_items: list):
    """根据 title+bank 从原始 items 中查找匹配项，用于补充字段"""
    llm_title = llm_item.get("title", "").strip()
    llm_bank = llm_item.get("bank", "").strip()

    # 1. 精确匹配 title+bank
    for orig in original_items:
        if orig.get("title", "").strip() == llm_title and orig.get("bank", "").strip() == llm_bank:
            return orig

    # 2. 仅 title 精确匹配
    for orig in original_items:
        if orig.get("title", "").strip() == llm_title:
            return orig

    # 3. 标题相似度匹配（阈值 0.6，容忍 LLM 微调标题）
    for orig in original_items:
        if _title_similarity(orig.get("title", ""), llm_title) >= 0.6:
            return orig

    # 4. bank 匹配 + 标题包含关系（LLM 可能缩短标题）
    if llm_bank:
        for orig in original_items:
            if orig.get("bank", "").strip() == llm_bank:
                orig_title = orig.get("title", "")
                if llm_title in orig_title or orig_title in llm_title:
                    return orig

    return None


def _fallback_merge(items: list, sources: list) -> dict:
    """LLM 不可用时的 fallback：使用标题相似度去重"""
    result = []

    for item in items:
        is_dup = False
        for i, existing in enumerate(result):
            if (_title_similarity(existing["title"], item["title"]) >= 0.85
                    and existing.get("bank", "") == item.get("bank", "")):
                # 保留内容更完整的
                existing_len = len(existing.get("raw_text", ""))
                new_len = len(item.get("raw_text", ""))
                if new_len > existing_len:
                    result[i] = item
                is_dup = True
                break
        if not is_dup:
            result.append(item)

    # 去重统计
    dropped_count = len(items) - len(result)
    print(f"  [Fallback] {len(items)} → {len(result)} items (dropped {dropped_count} duplicates)")

    batch_label = f"信用卡资讯（{len(sources)}份合并）"
    return {
        "batch_label": batch_label,
        "items": result,
        "stats": {
            "total_in": len(items),
            "total_out": len(result),
            "sources": sources,
            "merged": [],
            "dropped": [f"基础去重移除 {dropped_count} 条重复"],
        },
    }


# ═══════════════════════════════════════════════════════════════
#  convert_batch_to_merged — 标准 JSON → 旧 merged dict（供建议生成）
# ═══════════════════════════════════════════════════════════════


def convert_batch_to_merged(batch: dict) -> dict:
    """将标准 JSON batch 转为 generate_holistic_suggestion 所需的 merged dict"""
    # 按 category 分组构建 h1_sections
    categories = {}
    for item in batch.get("items", []):
        cat = item.get("category", "其他")
        if cat not in categories:
            categories[cat] = {
                "title": cat,
                "h2_items": [],
                "content": []
            }
        h2 = {
            "title": item.get("title", ""),
            "content": [item.get("raw_text", "")] if item.get("raw_text") else [],
            "images": [],
            "url": item.get("url", ""),
        }
        categories[cat]["h2_items"].append(h2)

    merged = {
        "title": batch.get("batch_label", "信用卡周报（整合版）"),
        "generated_at": datetime.now().isoformat(),
        "sources": batch.get("stats", {}).get("sources", []),
        "h1_sections": list(categories.values()),
    }
    return merged


# ═══════════════════════════════════════════════════════════════
#  持卡建议生成（保留原有逻辑）
# ═══════════════════════════════════════════════════════════════


def generate_holistic_suggestion(merged: dict) -> dict:
    """从整体汇总角度生成持卡建议（使用 LLM）"""
    full_text = ""
    try:
        # 汇总所有内容（使用层级结构）
        full_text_parts = []
        for h1 in merged.get("h1_sections", []):
            full_text_parts.append(f"【{h1['title']}】")
            for h2 in h1.get("h2_items", []):
                full_text_parts.append(f"  {h2['title']}")
                full_text_parts.extend(h2.get("content", [])[:5])
        full_text = "\n".join(full_text_parts)

        if len(full_text) < 50:
            return {"overall": None, "items": [], "error": "内容过少，无法生成建议"}

        # 尝试 LLM 整体分析
        sys.path.insert(0, str(PROJECT_ROOT / "card-holding-suggestion" / "scripts"))
        from scorer import load_config, DEFAULT_API_BASE, DEFAULT_MODEL

        cfg = load_config()
        api_key = cfg.get("api_key") or os.environ.get("LLM_API_KEY", "")
        api_base = cfg.get("api_base") or os.environ.get("LLM_API_BASE", DEFAULT_API_BASE)
        model = cfg.get("model") or os.environ.get("LLM_MODEL", DEFAULT_MODEL)

        # 如果可用，补充 RAG 检索上下文以提高建议准确性
        rag_context = None
        try:
            import rag_query
            entries = rag_query.load_kb()
            bm25 = rag_query.build_or_load_bm25(entries)
            scores = bm25.search(full_text, top_k=3)
            if scores:
                _, _, sources = rag_query.build_prompt(full_text[:200], entries, scores)
                rag_context = sources
        except Exception as e:
            print(f"  [Warn] RAG query failed (proceeding without RAG context): {e}")
            rag_context = None

        if api_key:
            prompt = f"""你是一个专业的信用卡分析师。请基于以下信用卡资讯汇总，从持卡人角度给出综合分析和建议。

【本期资讯汇总】
{full_text[:3000]}

【参考知识库检索（RAG）】
{json.dumps(rag_context, ensure_ascii=False) if rag_context else '无'}

### 分析要求：
1. 识别本期最重要的3-5条资讯（对持卡人影响最大的）
2. 从持卡人角度给出综合建议：
   - 哪些活动值得参与？为什么？
   - 哪些卡片值得关注/办理？
   - 有哪些风险需要注意？
   - 整体用卡策略建议
3. 按银行分类给出建议（如有明显差异）

### 输出格式（严格 JSON）：
```json
{{
  "highlights": [
    {{"title": "重要资讯标题", "bank": "银行", "reason": "重要性说明"}}
  ],
  "action_items": [
    {{"action": "建议动作", "priority": "高/中/低", "detail": "具体说明"}}
  ],
  "risk_warnings": ["风险提示1", "风险提示2"],
  "overall_strategy": "整体用卡策略建议（100字内）",
  "bank_summary": {{
    "银行名": "该银行本期要点总结"
  }}
}}
```"""

            from common.llm_client import call_llm_simple
            reply, llm_error = call_llm_simple(
                "", prompt,
                temperature=0.3, max_tokens=2048, timeout=120,
            )
            if llm_error or not reply:
                print(f"  [Warn] 建议 LLM 调用失败: {llm_error}")
                return _keyword_holistic_suggestion(full_text)

            # 解析 JSON
            m = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", reply, re.DOTALL)
            json_text = m.group(1).strip() if m else reply.strip()

            # 如果 code block 提取失败，尝试直接提取 JSON 对象
            if not json_text or not json_text.startswith("{"):
                m2 = re.search(r"\{[\s\S]*\}", reply)
                if m2:
                    json_text = m2.group()

            # 尝试多种方式解析 JSON
            overall = None

            # 方式1：直接解析
            try:
                overall = json.loads(json_text.strip())
            except json.JSONDecodeError:
                pass

            # 方式2：提取第一个完整的 JSON 对象
            if not overall:
                m2 = re.search(r"\{[\s\S]*\}", json_text)
                if m2:
                    try:
                        overall = json.loads(m2.group())
                    except json.JSONDecodeError:
                        # 尝试修复常见问题（尾部逗号等）
                        fixed = re.sub(r',\s*([}\]])', r'\1', m2.group())
                        try:
                            overall = json.loads(fixed)
                        except json.JSONDecodeError:
                            pass

            if overall:
                return {"overall": overall, "items": [], "generated_at": datetime.now().isoformat(), "scorer": "llm"}

        # 降级：关键词整体分析
        return _keyword_holistic_suggestion(full_text)

    except Exception as e:
        print(f"  [警告] LLM整体分析失败: {e}")
        return _keyword_holistic_suggestion(full_text)


def _keyword_holistic_suggestion(full_text: str) -> dict:
    """关键词降级的整体分析"""
    highlights = []
    action_items = []
    risk_warnings = []

    # 简单关键词匹配
    if "免年费" in full_text:
        action_items.append({"action": "关注免年费卡", "priority": "高", "detail": "有多款免年费卡可选"})
    if "返现" in full_text:
        action_items.append({"action": "参与返现活动", "priority": "中", "detail": "多个返现活动进行中"})
    if "缩水" in full_text or "取消" in full_text:
        risk_warnings.append("部分权益有缩水风险，请关注")
    if "升级" in full_text or "新增" in full_text:
        highlights.append({"title": "权益升级资讯", "bank": "多银行", "reason": "有新增或升级的权益"})

    overall = {
        "highlights": highlights[:5],
        "action_items": action_items[:5],
        "risk_warnings": risk_warnings[:3],
        "overall_strategy": "建议关注本期免年费卡和返现活动，注意权益变更风险。",
        "bank_summary": {}
    }
    return {"overall": overall, "items": [], "generated_at": datetime.now().isoformat(), "scorer": "keyword"}


# ═══════════════════════════════════════════════════════════════
#  主函数 — 新流程
# ═══════════════════════════════════════════════════════════════


def main():
    """主函数 — LLM 智能合并 + word-merger 生成 Word"""
    parser = argparse.ArgumentParser(description='多Word文档整合（LLM 智能合并）')
    parser.add_argument('--input', nargs='+', required=True, help='输入 docx 文件列表')
    parser.add_argument('--output', default='', help='输出文件路径（.docx）')
    parser.add_argument('--skip-analysis', action='store_true',
                        help='跳过持卡分析，仅输出 JSON + Markdown')
    parser.add_argument('--skip-images', action='store_true',
                        help='不嵌入图片到最终 Word')

    args = parser.parse_args()

    print("=" * 60)
    print("  Multi-Document Merge (LLM Smart)" + ("" if args.skip_analysis else " + Card Suggestions"))
    print("=" * 60)

    # ── 1. 读取所有文档 ──────────────────────────────────────
    print("\n[1/5] 读取文档...")
    contents = []
    for path in args.input:
        content = read_docx_content(path)
        h1_count = len(content.get("h1_sections", []))
        h2_count = sum(len(s.get("h2_items", [])) for s in content.get("h1_sections", []))
        img_count = len(content.get("images_map", {}))
        contents.append(content)
        print(f"  OK {Path(path).name}: {h1_count} categories, {h2_count} items, {img_count} images")

    # ── 2. 转为标准 JSON items ──────────────────────────────
    print("\n[2/5] 提取结构化条目...")
    items = contents_to_items(contents)
    sources = [Path(c["file"]).name for c in contents]
    print(f"  提取 {len(items)} 条来自 {len(sources)} 份文档")

    # ── 3. LLM 智能合并 ────────────────────────────────────
    print("\n[3/5] LLM 智能合并...")
    merged_batch = llm_merge(items, sources)

    # ── 4. 写入标准 JSON ────────────────────────────────────
    if not args.output:
        now = datetime.now()
        week_num = (now.day + now.replace(day=1).weekday()) // 7 + 1
        week_label = f"{now.year}年{now.month:02d}月第{week_num}周"
        args.output = str(PROJECT_ROOT / "data" / f"Merged_Report_{week_label}.docx")

    json_path = args.output.replace(".docx", ".json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(merged_batch, f, ensure_ascii=False, indent=2)
    print(f"  JSON: {Path(json_path).name}")

    # ── skip-analysis 路径：输出 JSON + MD 后结束 ───────────
    if args.skip_analysis:
        print("\n[Skip Analysis] --skip-analysis 已启用，跳过持卡分析与 Word 生成")
        # 输出 Markdown 预览
        md_path = args.output.replace(".docx", ".md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {merged_batch.get('batch_label', '信用卡周报')}\n\n")
            f.write(f"> 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
            f.write(f"**来源**: {', '.join(sources)}\n\n")
            f.write("---\n\n")
            for item in merged_batch.get("items", []):
                cat = item.get("category", "")
                title = item.get("title", "")
                bank = item.get("bank", "")
                f.write(f"## [{cat}] {title}\n\n")
                if bank:
                    f.write(f"**银行**: {bank}\n\n")
                summary = item.get("highlight_summary", "")
                if summary:
                    f.write(f"> {summary}\n\n")
                raw = item.get("raw_text", "")
                if raw:
                    for line in raw.split("\n")[:10]:
                        if line.strip():
                            f.write(f"{line}\n\n")
                if item.get("url"):
                    f.write(f"原文链接: {item['url']}\n\n")
                f.write("---\n\n")
            f.write(f"*由 merge_docs.py (LLM Smart) 生成*\n")
        print(f"  MD:   {Path(md_path).name}")
        print(f"\n[Success] Skip-analysis 完成")
        return True

    # ── 5a. 调用 word-merger/generate_report.py 生成 Word ───
    print(f"\n[4/5] 生成 Word: {Path(args.output).name}")
    generate_report_script = PROJECT_ROOT / "word-merger" / "scripts" / "generate_report.py"
    cmd = [sys.executable, "-X", "utf8", str(generate_report_script),
           "--input", json_path, "--output", args.output]
    if args.skip_images:
        cmd.append("--no-images")
    if merged_batch.get("batch_label"):
        cmd.extend(["--title", merged_batch["batch_label"]])

    result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace")
    if result.returncode != 0:
        print(f"  [Error] generate_report.py 失败:\n{result.stderr}")
        sys.exit(1)
    print(f"  Word 生成成功")

    # ── 5b. 生成持卡建议（可选） ─────────────────────────────
    print(f"\n[5/5] 生成持卡建议...")
    merged_for_suggestion = convert_batch_to_merged(merged_batch)
    suggestions = generate_holistic_suggestion(merged_for_suggestion)
    print(f"  建议: {'成功' if suggestions.get('overall') else '降级到关键词'}")

    # 将建议追加到 JSON
    merged_batch["suggestions"] = suggestions
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(merged_batch, f, ensure_ascii=False, indent=2)
    print(f"  JSON 已更新（含建议）")

    print(f"\n{'=' * 60}")
    print(f"  完成！")
    print(f"  Word: {args.output}")
    print(f"  JSON: {json_path}")
    print(f"{'=' * 60}")

    return True


if __name__ == "__main__":
    main()
