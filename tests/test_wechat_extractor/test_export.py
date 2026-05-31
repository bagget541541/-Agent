"""wechat-article-extractor / export_document.py 单元测试"""

import os
import re
import sys
import tempfile
from pathlib import Path

import pytest

# ── 路径设置 ──────────────────────────────────────────────────
_PROJECT_ROOT = str(Path(__file__).resolve().parents[2])
sys.path.insert(0, _PROJECT_ROOT)
sys.path.insert(0, os.path.join(_PROJECT_ROOT, "wechat-article-extractor", "scripts"))

import export_document as ed


# ═══════════════════════════════════════════════════════════════
# 正则常量
# ═══════════════════════════════════════════════════════════════

class TestImagePattern:

    def test_basic_match(self):
        m = ed.IMAGE_PATTERN.search("![描述](images/pic.png)")
        assert m is not None
        assert m.group(1) == "描述"
        assert m.group(2) == "images/pic.png"

    def test_empty_alt_text(self):
        m = ed.IMAGE_PATTERN.search("![](data/img.jpg)")
        assert m is not None
        assert m.group(1) == ""

    def test_no_match_on_plain_text(self):
        m = ed.IMAGE_PATTERN.search("这是一段普通文字，没有图片")
        assert m is None

    def test_url_with_query_params(self):
        url = "https://img.example.com/a.jpg?token=abc&size=large"
        m = ed.IMAGE_PATTERN.search(f"![]({url})")
        assert m is not None
        assert m.group(2) == url


class TestHeadingRe:

    @pytest.mark.parametrize("text, level, content", [
        ("# 一级标题", 1, "一级标题"),
        ("## 二级标题", 2, "二级标题"),
        ("### 三级标题", 3, "三级标题"),
        ("#### 四级", 4, "四级"),
        ("##### 五级", 5, "五级"),
        ("###### 六级", 6, "六级"),
    ])
    def test_heading_levels(self, text, level, content):
        m = ed.HEADING_RE.match(text)
        assert m is not None
        assert len(m.group(1)) == level
        assert m.group(2) == content

    def test_no_match_on_plain_text(self):
        assert ed.HEADING_RE.match("这不是标题") is None

    def test_no_match_on_single_hash_no_space(self):
        assert ed.HEADING_RE.match("#无空格") is None

    def test_no_match_on_seven_hashes(self):
        assert ed.HEADING_RE.match("####### 太多了") is None


# ═══════════════════════════════════════════════════════════════
# export_markdown — 字面量 \n 处理
# ═══════════════════════════════════════════════════════════════

class TestExportMarkdown:

    def test_literal_backslash_n_split(self):
        """含字面量 \\n 的内容应按行分割后写入。"""
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.md")
            content = "第一行\\n第二行\\n第三行"
            ed.export_markdown(content, out)
            with open(out, "r", encoding="utf-8") as f:
                result = f.read()
            lines = result.strip().split("\n")
            assert lines == ["第一行", "第二行", "第三行"]

    def test_no_literal_n_write_directly(self):
        """不含字面量 \\n 时直接写入原文。"""
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.md")
            content = "正常内容\n含真换行"
            ed.export_markdown(content, out)
            with open(out, "r", encoding="utf-8") as f:
                result = f.read()
            assert result == content

    def test_empty_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.md")
            ed.export_markdown("", out)
            with open(out, "r", encoding="utf-8") as f:
                assert f.read() == ""

    def test_returns_output_path(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.md")
            result = ed.export_markdown("hello", out)
            assert result == out


# ═══════════════════════════════════════════════════════════════
# export_word — 结构化文档生成
# ═══════════════════════════════════════════════════════════════

class TestExportWord:

    def test_heading_level_1(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.docx")
            ed.export_word("# 一级标题", out)
            from docx import Document
            doc = Document(out)
            headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
            assert len(headings) == 1
            assert headings[0].text == "一级标题"

    def test_heading_level_2(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.docx")
            ed.export_word("## 二级标题", out)
            from docx import Document
            doc = Document(out)
            headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
            assert len(headings) == 1
            assert headings[0].text == "二级标题"

    def test_field_name_bolded(self):
        """已知字段名 + 冒号 → 字段名部分加粗。"""
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.docx")
            ed.export_word("活动内容：618大促满200减50", out)
            from docx import Document
            doc = Document(out)
            # 找到包含 '活动内容' 的段落
            target = [p for p in doc.paragraphs if "活动内容" in p.text]
            assert len(target) == 1
            runs = target[0].runs
            # 第一个 run 应该是加粗的
            assert runs[0].bold is True
            assert "活动内容" in runs[0].text

    def test_field_name_with_colon_space(self):
        """字段名 + 中间有空格的冒号。"""
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.docx")
            ed.export_word("卡种 ： 经典白金卡", out)
            from docx import Document
            doc = Document(out)
            target = [p for p in doc.paragraphs if "卡种" in p.text]
            assert len(target) == 1
            assert target[0].runs[0].bold is True

    def test_plain_text_paragraph(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.docx")
            ed.export_word("这是一段普通文字", out)
            from docx import Document
            doc = Document(out)
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            assert len(paragraphs) == 1
            assert paragraphs[0].text == "这是一段普通文字"

    def test_image_inserted_when_file_exists(self):
        with tempfile.TemporaryDirectory() as tmp:
            # 创建一张 1x1 像素的 PNG 图片
            img_path = os.path.join(tmp, "test.png")
            from PIL import Image
            Image.new("RGB", (10, 10), color="blue").save(img_path)

            out = os.path.join(tmp, "test.docx")
            ed.export_word("![测试图片](test.png)", out)

            from docx import Document
            doc = Document(out)
            # 检查是否有图片被插入（inline shapes）
            img_count = sum(1 for p in doc.paragraphs
                           for r in p.runs
                           if r._r.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))
            # 至少有一个段落包含 drawing 元素
            assert img_count >= 1

    def test_image_not_found_placeholder(self):
        """图片文件不存在时插入占位文字。"""
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.docx")
            ed.export_word("![缺少图片](nonexistent.png)", out)
            from docx import Document
            doc = Document(out)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "不存在" in all_text or "图片" in all_text

    def test_multiple_images_in_line(self):
        """单行中包含多个图片标记。"""
        with tempfile.TemporaryDirectory() as tmp:
            img1 = os.path.join(tmp, "a.png")
            img2 = os.path.join(tmp, "b.png")
            from PIL import Image
            Image.new("RGB", (10, 10)).save(img1)
            Image.new("RGB", (10, 10)).save(img2)

            out = os.path.join(tmp, "test.docx")
            ed.export_word("![a](a.png) 和 ![b](b.png)", out)
            from docx import Document
            doc = Document(out)
            # 应该有 a、"和"、b 三个段落/内容
            assert len(doc.paragraphs) >= 2

    def test_returns_output_path(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = os.path.join(tmp, "test.docx")
            result = ed.export_word("hello", out)
            assert result == out
