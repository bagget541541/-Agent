#!/usr/bin/env python3
"""
Word文档读取器
读取多个Word文档并提取结构化文本内容、表格和图片
"""

import sys
import os
import argparse
import json
import hashlib
from docx import Document
from docx.oxml.ns import qn


def get_paragraph_style(para):
    """
    获取段落的样式类型
    """
    style = para.style
    if style is None:
        return {"style_type": "Normal", "style_name": "Normal"}

    style_name = style.name or "Normal"

    if style_name.startswith("Heading"):
        return {"style_type": "Heading", "style_name": style_name}
    elif "List" in style_name or "Numbering" in style_name:
        if "Bullet" in style_name:
            return {"style_type": "ListBullet", "style_name": style_name}
        elif "Number" in style_name:
            return {"style_type": "ListNumber", "style_name": style_name}
        else:
            return {"style_type": "ListBullet", "style_name": style_name}
    elif "Quote" in style_name or "Citation" in style_name:
        return {"style_type": "Quote", "style_name": style_name}
    else:
        return {"style_type": "Normal", "style_name": style_name}


def extract_images_from_paragraph(para, doc):
    """
    从段落中提取图片信息

    Returns:
        list: 图片信息列表，包含图片ID和位置
    """
    images = []
    for run in para.runs:
        # 检查run中是否有图片
        for elem in run._element:
            # 检查是否是drawing元素（包含图片）
            if elem.tag.endswith('}drawing') or elem.tag.endswith('}pict'):
                # 获取底层的blip图片
                for blip in elem.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed and embed in doc.part.rels:
                        rel = doc.part.rels[embed]
                        if "image" in rel.reltype:
                            images.append({
                                "image_id": embed,
                                "source_ref": rel.target_ref
                            })
    return images


def extract_all_images(doc):
    """
    提取文档中所有图片

    Returns:
        dict: 图片ID到文件路径的映射
    """
    images_map = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob

                # 生成唯一文件名
                image_hash = hashlib.md5(image_bytes).hexdigest()[:8]
                ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                if not ext:
                    ext = 'png'
                file_name = f"image_{rel_id}_{image_hash}.{ext}"

                images_map[rel_id] = {
                    "file_name": file_name,
                    "ext": ext,
                    "source_ref": rel.target_ref,
                    "size": len(image_bytes)
                }
            except Exception as e:
                continue
    return images_map


def extract_paragraphs_with_images(doc, images_map):
    """
    提取所有段落，标记图片位置

    Returns:
        list: 段落列表，包含文本、样式和图片信息
    """
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip() or len(para.runs) > 0:
            style_info = get_paragraph_style(para)

            # 检查段落中的图片
            images = []
            for run in para.runs:
                for elem in run._element.iter():
                    if 'blip' in elem.tag:
                        embed = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in images_map:
                            images.append({
                                "image_id": embed,
                                "file_name": images_map[embed]["file_name"]
                            })

            paragraphs.append({
                "text": para.text.strip(),
                "style_type": style_info["style_type"],
                "style_name": style_info["style_name"],
                "images": images
            })
    return paragraphs


def extract_tables(doc):
    """
    提取所有表格（包含图片）
    """
    tables_data = []
    for table_idx, table in enumerate(doc.tables):
        table_info = {
            "table_index": table_idx,
            "rows": []
        }
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                # 检查单元格中的图片
                cell_images = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        for elem in run._element.iter():
                            if 'blip' in elem.tag:
                                embed = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed and embed in images_map:
                                    cell_images.append({
                                        "image_id": embed,
                                        "file_name": images_map[embed]["file_name"]
                                    })

                cell_info = {"text": cell_text, "images": cell_images}
                row_cells.append(cell_info)

            if any(c["text"] or c["images"] for c in row_cells):
                table_info["rows"].append({
                    "row_index": row_idx,
                    "cells": row_cells
                })

        if table_info["rows"]:
            tables_data.append(table_info)
    return tables_data


def extract_images_from_tables(doc, images_map):
    """从表格单元格中提取图片"""
    tables_with_images = []
    for table_idx, table in enumerate(doc.tables):
        table_images = []
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for elem in run._element.iter():
                            if 'blip' in elem.tag:
                                embed = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed and embed in images_map:
                                    table_images.append({
                                        "image_id": embed,
                                        "file_name": images_map[embed]["file_name"]
                                    })
        if table_images:
            tables_with_images.append({
                "table_index": table_idx,
                "images": table_images
            })
    return tables_with_images


def read_docx(file_path, output_dir=None):
    """
    读取单个Word文档

    Args:
        file_path: Word文档路径
        output_dir: 图片输出目录，若提供则保存图片

    Returns:
        dict: 包含文档名、段落（带样式）、表格和图片信息的字典
    """
    try:
        doc = Document(file_path)

        # 提取所有图片
        images_map = extract_all_images(doc)

        # 如果提供了输出目录，保存图片
        saved_images = []
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            for rel_id, img_info in images_map.items():
                try:
                    image_part = doc.part.rels[rel_id].target_part
                    image_bytes = image_part.blob
                    output_path = os.path.join(output_dir, img_info["file_name"])
                    with open(output_path, 'wb') as f:
                        f.write(image_bytes)
                    saved_images.append({
                        "image_id": rel_id,
                        "file_name": img_info["file_name"],
                        "saved_path": output_path
                    })
                except Exception as e:
                    continue

        # 提取段落（带图片信息）
        paragraphs = extract_paragraphs_with_images(doc, images_map)

        # 提取表格
        tables = extract_tables(doc)

        # 统计各类型段落数量
        style_counts = {}
        for p in paragraphs:
            stype = p["style_type"]
            style_counts[stype] = style_counts.get(stype, 0) + 1

        # 统计图片数量
        image_count = len(images_map)

        return {
            "file_name": file_path,
            "paragraphs": paragraphs,
            "tables": tables,
            "images": {
                "total": image_count,
                "list": [{"image_id": k, **v} for k, v in images_map.items()],
                "saved": saved_images if output_dir else []
            },
            "style_counts": style_counts,
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables),
            "status": "success"
        }
    except FileNotFoundError:
        return {
            "file_name": file_path,
            "status": "error",
            "error": "文件不存在"
        }
    except Exception as e:
        return {
            "file_name": file_path,
            "status": "error",
            "error": str(e)
        }


def main():
    parser = argparse.ArgumentParser(description='读取Word文档并提取结构化文本内容')
    parser.add_argument('--input-files', required=True,
                        help='输入的Word文档路径，多个文件用逗号分隔')
    parser.add_argument('--output', default=None,
                        help='可选：指定输出文件路径，将JSON结果写入文件')
    parser.add_argument('--images-dir', default=None,
                        help='可选：指定图片输出目录，图片将保存到此目录')

    args = parser.parse_args()

    # 解析文件列表
    file_paths = [f.strip() for f in args.input_files.split(',') if f.strip()]

    if not file_paths:
        error_msg = json.dumps({"error": "未提供有效的文件路径"})
        print(error_msg)
        sys.exit(1)

    # 读取所有文档
    results = []
    for file_path in file_paths:
        # 为每个文档创建单独的子目录
        doc_images_dir = None
        if args.images_dir:
            doc_name = os.path.splitext(os.path.basename(file_path))[0]
            doc_images_dir = os.path.join(args.images_dir, doc_name)
        result = read_docx(file_path, doc_images_dir)
        results.append(result)

    # 构建输出
    output = {
        "total_files": len(file_paths),
        "success_count": sum(1 for r in results if r.get("status") == "success"),
        "images_base_dir": args.images_dir,
        "files": results
    }

    json_output = json.dumps(output, ensure_ascii=False, indent=2)

    # 根据参数决定输出方式
    if args.output:
        try:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(json_output)
            print(json.dumps({
                "status": "success",
                "output_file": args.output,
                "files_processed": len(results),
                "images_dir": args.images_dir
            }, ensure_ascii=False))
        except Exception as e:
            print(json.dumps({
                "status": "error",
                "error": f"写入文件失败: {str(e)}"
            }))
            sys.exit(1)
    else:
        print(json_output)


if __name__ == "__main__":
    main()
