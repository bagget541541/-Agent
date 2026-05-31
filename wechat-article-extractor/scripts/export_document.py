#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档导出脚本

功能：将格式化内容导出为 Word 或 Markdown 文档
"""

import sys
import argparse
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os, sys

# 将 word-merger utils 加入路径以共享 IMAGE_WIDTH_CM 常量
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'word-merger', 'scripts')))
try:
    from utils import IMAGE_WIDTH_CM
except Exception:
    IMAGE_WIDTH_CM = 13

IMAGE_PATTERN = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

# 已知字段名集合（冒号结尾时加粗，但不作为顶级标题）
KNOWN_FIELD_NAMES = {
    '活动内容', '活动时间', '适用人群',
    '卡种', '卡亮点', '详情', '来源',
    '消息时间', '影响范围', '变更内容', '变更分析',
    '消息内容', '点评',
    '卡组织', '发卡银行', '年费标准', '有效期限',
    '原文链接', '发布时间', '作者',
}

# Markdown 标题正则（支持 # 和 ##）
HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$')


def export_word(content, output_path):
    """
    导出为 Word 文档

    Args:
        content: 格式化的文本内容（包含Markdown图片语法）
        output_path: 输出文件路径
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    # 分割内容，逐行处理
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue

        # ── 1) 检查 Markdown 标题语法（# 开头） ──
        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))  # # = 1, ## = 2 ...
            text = heading_match.group(2)
            heading = doc.add_heading(text, level=level)
            # 确保标题也使用宋体
            for run in heading.runs:
                run.font.name = '宋体'
            continue

        # ── 2) 检查是否包含图片标记 ──
        img_matches = list(IMAGE_PATTERN.finditer(line))

        if img_matches:
            # 行中包含图片，需要分割文本和图片
            last_end = 0
            for match in img_matches:
                # 添加图片前的文本
                if match.start() > last_end:
                    text_part = line[last_end:match.start()]
                    if text_part.strip():
                        doc.add_paragraph(text_part.strip())

                # 添加图片
                img_path = match.group(2)
                # 支持相对路径和绝对路径
                original_path = img_path  # 保存原始路径用于调试
                if not os.path.isabs(img_path):
                    # 相对路径，基于输出文件所在目录解析
                    if output_path:
                        base_dir = os.path.dirname(os.path.abspath(output_path))
                        img_path = os.path.join(base_dir, img_path)
                    else:
                        img_path = os.path.abspath(img_path)

                # 检查图片文件是否存在
                if os.path.exists(img_path):
                    try:
                        # 添加图片，设置宽度与 word-merger 统一
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(img_path, width=Cm(IMAGE_WIDTH_CM))
                        # 图片居中
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception as e:
                        # 图片添加失败，添加占位文本
                        print(f"[警告] 图片加载失败: {original_path} -> {img_path}, 错误: {str(e)}", file=sys.stderr)
                        doc.add_paragraph(f"[图片加载失败: {os.path.basename(img_path)}]")
                else:
                    # 图片文件不存在
                    print(f"[警告] 图片文件不存在: {original_path} -> {img_path}", file=sys.stderr)
                    doc.add_paragraph(f"[图片不存在: {os.path.basename(img_path)}]")

                last_end = match.end()

            # 添加图片后的文本
            if last_end < len(line):
                text_part = line[last_end:]
                if text_part.strip():
                    doc.add_paragraph(text_part.strip())
        else:
            # 没有图片，正常处理文本
            # 判断是否以已知字段名开头接冒号（如"活动内容：xxx"）
            field_match = re.match(
                r'^\s*(' + '|'.join(re.escape(n) for n in sorted(KNOWN_FIELD_NAMES, key=len, reverse=True)) + r')\s*[：:]\s*(.*)$',
                line
            )
            if field_match:
                field_name = field_match.group(1)
                field_value = field_match.group(2)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(field_name + '：')
                run.bold = True
                run.font.size = Pt(12)
                if field_value.strip():
                    run2 = paragraph.add_run(field_value)
                    run2.font.size = Pt(12)
            else:
                doc.add_paragraph(line)

    doc.save(output_path)
    return output_path


def export_markdown(content, output_path):
    """
    导出为 Markdown 文档

    Args:
        content: 格式化的文本内容
        output_path: 输出文件路径
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        # 检查内容中是否包含字面量 \n 字符（转义的换行符）
        if '\\n' in content:
            # 如果包含 \n 字符，将其转换为实际换行符
            lines = content.split('\\n')
            for line in lines:
                f.write(line + '\n')
        else:
            # 否则直接写入内容
            f.write(content)

    return output_path


def main():
    parser = argparse.ArgumentParser(description='导出文档')
    parser.add_argument('--content', required=False, help='格式化的文本内容')
    parser.add_argument('--content-file', required=False, help='包含内容的文件路径（与--content二选一）')
    parser.add_argument('--format', required=True, choices=['word', 'md'], help='输出格式')
    parser.add_argument('--output', required=True, help='输出文件路径')

    args = parser.parse_args()

    # 检查参数
    if not args.content and not args.content_file:
        print("错误: 必须提供 --content 或 --content-file 参数", file=sys.stderr)
        sys.exit(1)

    if args.content and args.content_file:
        print("错误: --content 和 --content-file 不能同时使用", file=sys.stderr)
        sys.exit(1)

    # 获取内容
    if args.content_file:
        # 从文件读取内容
        try:
            with open(args.content_file, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            print(f"错误: 无法读取文件 {args.content_file}: {str(e)}", file=sys.stderr)
            sys.exit(1)
    else:
        content = args.content

    try:
        if args.format == 'word':
            output_path = export_word(content, args.output)
        else:
            output_path = export_markdown(content, args.output)

        print(f"文档已成功导出: {output_path}")
    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
