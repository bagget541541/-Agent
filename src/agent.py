#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
信用卡周报全流程 Agent

全流程编排入口，串联所有 skill：
  Step 1: 抓取微信文章
  Step 2: 抓取银行公告
  Step 3: 分类 + 合并
  Step 4: 生成 Word 周报
  Step 5: 持卡用卡分析
  Step 6: 追加持卡建议到周报
  Step 7: 归档到知识库
"""

import argparse
import json
import os
import subprocess
import sys
import time
from datetime import datetime, timedelta
from pathlib import Path
import tempfile
import shutil

# 添加项目根目录到路径
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from common.schema import CreditCardBatch, CreditCardItem
from common.utils import get_week_label
from common.normalizer import normalize_item, normalize_topic
from common.review import export_review_queue
from common.article_envelope import build_article_envelope
from common.topic_splitter import detect_multi_topic, split_article_into_topics, detect_multi_topic_from_raw_text
from rag_query import query_for_suggestions


def _item_to_mapping(item):
    """兼容 dict / CreditCardItem / 其他带属性对象。"""
    if isinstance(item, dict):
        return item
    if hasattr(item, "to_dict") and callable(getattr(item, "to_dict")):
        return item.to_dict()
    return {
        "source": getattr(item, "source", ""),
        "category": getattr(item, "category", ""),
        "bank": getattr(item, "bank", ""),
        "title": getattr(item, "title", ""),
        "url": getattr(item, "url", ""),
        "raw_text": getattr(item, "raw_text", ""),
        "images": getattr(item, "images", []) or [],
        "structured": getattr(item, "structured", {}) or {},
        "author": getattr(item, "author", ""),
        "publish_time": getattr(item, "publish_time", ""),
        "extracted_at": getattr(item, "extracted_at", ""),
    }


def _build_item_from_raw_text_topic(topic: dict, source_article: dict) -> dict:
    """从 raw_text 主题拆分结果构造可被 normalize_item 消费的 dict。"""
    bank_hints = topic.get("bank_hints", [])
    first_bank = bank_hints[0] if bank_hints else ""
    return {
        "title": topic.get("headline", ""),
        "raw_text": topic.get("body_text", ""),
        "content_text": topic.get("body_text", ""),
        "url": source_article.get("url", ""),
        "account_name": source_article.get("account_name", "") or source_article.get("author", "") or "未知公众号",
        "images": source_article.get("images", []),
        "source": "wechat",
        "bank": first_bank,
        "publish_time": source_article.get("fetched_at", "") or source_article.get("publish_time", ""),
        "content_blocks": [],
    }


def _load_json_fallback(filepath: Path):
    """尽量用多种编码读取 JSON。"""
    for encoding in ("utf-8", "utf-8-sig", "gbk", "cp936"):
        try:
            with open(filepath, "r", encoding=encoding) as f:
                return json.load(f)
        except Exception:
            continue
    return None


def _safe_write_json(filepath: Path, data: dict, *, ensure_ascii: bool = False, indent: int = 2):
    """安全写入 JSON：优先尝试 portalocker 加文件锁，否则使用原子临时文件替换。"""
    try:
        import portalocker
    except Exception:
        portalocker = None

    tmp_path = filepath.with_suffix(filepath.suffix + ".tmp")
    # 确保目录存在
    os.makedirs(filepath.parent, exist_ok=True)
    try:
        # 写到临时文件
        with open(tmp_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=ensure_ascii, indent=indent)

        if portalocker:
            # 再用锁替换目标文件
            with open(str(filepath), 'a+', encoding='utf-8') as dest_f:
                try:
                    portalocker.lock(dest_f, portalocker.LOCK_EX)
                    # 替换文件
                    dest_f.close()
                    os.replace(tmp_path, filepath)
                finally:
                    try:
                        portalocker.unlock(dest_f)
                    except Exception:
                        pass
        else:
            # 原子替换
            os.replace(tmp_path, filepath)
    finally:
        # 清理残留 tmp 文件
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


def _normalize_wechat_article(raw: dict, url: str = "") -> dict:
    """把缓存/抓取结果统一成 step1 输出格式。"""
    article_url = raw.get("url") or url
    title = raw.get("title") or raw.get("original_title") or raw.get("name") or "无标题"
    content_text = (
        raw.get("content_text")
        or raw.get("text")
        or raw.get("full_text")
        or raw.get("summary")
        or ""
    )
    images = raw.get("images") or []
    # 正文为空但有图片时，用标题做文本兜底
    if not content_text and images:
        content_text = title
    # 无论来源如何，每个条目最多保留 3 张图片（前 3 张通常是标题卡+关键内容）
    if len(images) > 3:
        images = images[:3]
    account_name = raw.get("account_name") or raw.get("author") or raw.get("biz_name") or "未知公众号"

    # 保留 content_blocks（用于多主题拆分）
    content_blocks = raw.get("content_blocks") or raw.get("blocks") or []

    return {
        "title": title,
        "url": article_url,
        "account_name": account_name,
        "content_text": content_text,
        "content_blocks": content_blocks,
        "images": images,
        "source": raw.get("source", "wechat-cache"),
        "fetched_at": raw.get("fetched_at", datetime.now().isoformat()),
    }


def _load_cached_wechat_articles() -> list[dict]:
    """从本地缓存加载公众号文章，作为无 URL 时的回退来源。"""
    cached_articles: list[dict] = []

    monitor_dir = PROJECT_ROOT / "data" / "wechat_monitor" / "articles"
    if monitor_dir.exists():
        for cache_file in sorted(monitor_dir.glob("*.json")):
            data = _load_json_fallback(cache_file)
            if not data:
                continue
            if isinstance(data, list):
                for item in data:
                    if isinstance(item, dict):
                        cached_articles.append(_normalize_wechat_article(item))
            elif isinstance(data, dict):
                # 既兼容 {"url": {...}}，也兼容单篇文章字典
                if data and all(isinstance(value, dict) for value in data.values()) and any(
                    isinstance(key, str) and key.startswith("http") for key in data.keys()
                ):
                    for url, item in data.items():
                        cached_articles.append(_normalize_wechat_article(item, url=url))
                else:
                    cached_articles.append(_normalize_wechat_article(data))

    if cached_articles:
        return cached_articles

    legacy_path = PROJECT_ROOT / "data" / "wechat_articles.json"
    if legacy_path.exists():
        data = _load_json_fallback(legacy_path)
        if isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    cached_articles.append(_normalize_wechat_article(item))
        elif isinstance(data, dict):
            for url, item in data.items():
                if isinstance(item, dict):
                    cached_articles.append(_normalize_wechat_article(item, url=url))

    return cached_articles


def _analysis_error(error_code: str, message: str, timeout_seconds: int = None, scorer: str = "", model: str = "") -> dict:
    """构造可追踪的分析失败结果。"""
    payload = {
        "status": "error",
        "error_code": error_code,
        "error_message": message,
    }
    if timeout_seconds is not None:
        payload["timeout_seconds"] = timeout_seconds
    if scorer:
        payload["scorer"] = scorer
    if model:
        payload["model"] = model
    return payload


def step1_fetch_wechat(urls: list[str] = None) -> list[dict]:
    """
    Step 1: 抓取微信文章

    使用 wechat-article-extractor（requests + Playwright 兜底）
    保留 URL 去重机制

    Args:
        urls: 微信文章 URL 列表，None 则从配置读取

    Returns:
        文章列表（dict 格式）
    """
    print("\n" + "=" * 60)
    print("Step 1: Fetch WeChat articles")
    print("=" * 60)

    cached_articles: list[dict] = _load_cached_wechat_articles()

    def _needs_refetch(article: dict) -> bool:
        """是否需要重抓。检查文本完整性及图片文件是否存在。"""
        # 检查文本是否为空
        txt = (article.get("content_text") or article.get("text") or article.get("full_text") or "")
        if not txt or str(txt).strip() == "":
            return True
        # 检查图片文件是否真实存在
        images = article.get("images") or []
        if images:
            for img_path in images:
                # 本地路径：直接检查文件存在
                if os.path.exists(img_path):
                    continue
                # HTTP 路径：已经被下载到本地但存储路径可能是映射后路径
                # 如果全是远程 URL 则视为有效（后续会下载）
                if img_path.startswith("http://") or img_path.startswith("https://"):
                    continue
                return True  # 存在本地路径但文件不存在，需要重抓
        return False

    def _attempt_refetch(urls_to_refetch: list[str]) -> list[dict]:
        """尝试对给定 URL 列表重新抓取文章，返回标准化后的文章列表（只包含成功的）。"""
        if not urls_to_refetch:
            return []
        extractor_path = PROJECT_ROOT / "wechat-article-extractor" / "scripts"
        _extractor_str = str(extractor_path)
        if _extractor_str not in sys.path:
            sys.path.insert(0, _extractor_str)
        try:
            from fetch_wechat_article import process_single
        except Exception as e:
            print(f"  [Refetch] import failed: {e}")
            return []

        refreshed = []
        images_dir = PROJECT_ROOT / "data" / "images"
        for url in urls_to_refetch:
            try:
                # 为了避免重复的 Processing 日志，直接一次性带图片抓取（兜底更全面）；限 3 张关键图
                art = process_single(url, download_images=True, images_dir=str(images_dir), extract_ocr=False, extract_llm=True, max_images=3)

                # 判定成功时检查多字段（content_text/text/full_text）
                text_candidates = (
                    (art.get('content_text') if isinstance(art, dict) else None) or
                    (art.get('text') if isinstance(art, dict) else None) or
                    (art.get('full_text') if isinstance(art, dict) else None) or
                    ""
                )

                if art and not art.get("error"):
                    # 有文本或有图片都算成功（兼容全图片文章）
                    has_text = bool(str(text_candidates).strip())
                    has_images = bool(art.get("images")) if isinstance(art, dict) else False
                    if has_text or has_images:
                        refreshed.append(_normalize_wechat_article(art, url=url))
                        print(f"  [Refetch OK] {art.get('title','')[:40]}...")
                    else:
                        print(f"  [Refetch Failed] {url}")

            except Exception as e:
                print(f"  [Refetch Exception] {url[:60]}...: {e}")

        return refreshed

    if not urls:
        if cached_articles:
            print(f"  [Cache] Loaded {len(cached_articles)} cached WeChat articles")
            # 检查缓存中文章是否缺失正文，尝试重抓
            urls_to_refetch = [a.get("url") for a in cached_articles if a.get("url") and _needs_refetch(a)]
            if urls_to_refetch:
                print(f"  [Refetch] {len(urls_to_refetch)} cached articles missing content, attempting refetch")
                refreshed = _attempt_refetch(urls_to_refetch)
                # 用重抓成功的条目替换缓存中的对应条目
                if refreshed:
                    by_url = {a['url']: a for a in refreshed}
                    new_list = []
                    for a in cached_articles:
                        if a.get('url') in by_url:
                            new_list.append(by_url[a.get('url')])
                        else:
                            new_list.append(a)
                    cached_articles = new_list

            return cached_articles

        print("[Skip] No WeChat article URLs provided and no cached articles found")
        return []

    # URL 去重（使用 wechat-rss-monitor 的历史记录）
    history_path = PROJECT_ROOT / "data" / "wechat_monitor" / "history.json"
    history = {"fetched_urls": []}
    if history_path.exists():
        try:
            with open(history_path, 'r', encoding='utf-8') as f:
                history = json.load(f)
        except Exception:
            pass

    # 过滤已抓取的 URL
    import hashlib
    new_urls = []
    for url in urls:
        url_hash = hashlib.md5(url.encode('utf-8')).hexdigest()
        if url_hash not in history.get("fetched_urls", []):
            new_urls.append(url)

    # 导入 wechat-article-extractor（提前导入，cache 为空时 re-fetch 也需要）
    extractor_path = PROJECT_ROOT / "wechat-article-extractor" / "scripts"
    _extractor_str = str(extractor_path)
    if _extractor_str not in sys.path:
        sys.path.insert(0, _extractor_str)

    try:
        from fetch_wechat_article import process_single
    except ImportError as e:
        print(f"  [Error] wechat-article-extractor import failed: {e}")
        return []

    articles = []
    if not new_urls:
        if cached_articles:
            print(f"  All URLs already fetched, using {len(cached_articles)} cached WeChat articles")
            # 仅保留匹配输入 URL 的条目（避免返回缓存中其他无关文章）
            matched = [a for a in cached_articles if a.get("url") in urls] if urls else cached_articles
            if matched:
                # 检查并尝试重抓正文缺失的缓存条目
                urls_to_refetch = [a.get("url") for a in matched if a.get("url") and _needs_refetch(a)]
                if urls_to_refetch:
                    print(f"  [Refetch] {len(urls_to_refetch)} cached articles missing content, attempting refetch")
                    refreshed = _attempt_refetch(urls_to_refetch)
                    if refreshed:
                        by_url = {a['url']: a for a in refreshed}
                        new_list = []
                        for a in matched:
                            if a.get('url') in by_url:
                                new_list.append(by_url[a.get('url')])
                            else:
                                new_list.append(a)
                        matched = new_list
                return matched
            # matched 为空（缓存文件存在但不含输入 URL）→ 走 re-fetch
            print(f"  [Warning] None of {len(urls)} input URLs found in cache, re-fetching")
        else:
            # cache 为空 —— history 记录了 hash 但缓存文件没有对应条目，强制重新抓取
            print(f"  [Warning] All {len(urls)} URL(s) marked fetched but cache empty, re-fetching")
        for url in urls:
            try:
                article = process_single(url, download_images=False)
                if article and not article.get("error"):
                    standard_article = {
                        "title": article.get("title") or article.get("original_title") or "无标题",
                        "url": url,
                        "account_name": article.get("account_name", "") or article.get("author", "") or "未知公众号",
                        "content_text": article.get("content_text", "") or article.get("text", "") or article.get("full_text", ""),
                        "images": article.get("images", []),
                        "source": article.get("source", "unknown"),
                        "fetched_at": datetime.now().isoformat(),
                    }
                    articles.append(standard_article)
                    print(f"  [OK] {standard_article['title'][:40]}...")
                time.sleep(1)
            except Exception as e:
                print(f"  [Failed] {url[:50]}...: {e}")
        return articles

    print(f"  To fetch: {len(new_urls)} articles ({len(urls) - len(new_urls)} already fetched)")

    # 抓取文章
    for url in new_urls:
        try:
            article = process_single(url, download_images=False)

            if article and not article.get("error"):
                # 转换为标准格式
                standard_article = {
                    "title": article.get("title") or article.get("original_title") or "无标题",
                    "url": url,
                    "account_name": article.get("account_name", "") or article.get("author", "") or "未知公众号",
                    "content_text": article.get("content_text", "") or article.get("text", "") or article.get("full_text", ""),
                    "images": article.get("images", []),
                    "source": article.get("source", "unknown"),
                    "fetched_at": datetime.now().isoformat(),
                }
                articles.append(standard_article)

                # Mark as fetched
                url_hash = hashlib.md5(url.encode('utf-8')).hexdigest()
                if "fetched_urls" not in history:
                    history["fetched_urls"] = []
                history["fetched_urls"].append(url_hash)

                print(f"  [OK] {standard_article['title'][:40]}...")

            time.sleep(1)  # Rate limiting

        except Exception as e:
            print(f"  [Failed] {url[:50]}...: {e}")

    # 对新抓取结果中正文为空的条目尝试重抓一次（带图片兜底）
    urls_missing = [a['url'] for a in articles if not (a.get('content_text') or '').strip()]
    if urls_missing:
        print(f"  [Refetch] {len(urls_missing)} fetched articles missing content, attempting refetch")
        refreshed = _attempt_refetch(urls_missing)
        if refreshed:
            by_url = {a['url']: a for a in refreshed}
            new_articles = []
            for a in articles:
                refreshed_art = by_url.get(a.get('url'))
                if refreshed_art:
                    # 重抓成功：有文本或有本地图片都替换
                    has_text = bool((refreshed_art.get('content_text') or '').strip())
                    has_images = bool(refreshed_art.get('images'))
                    if has_text or has_images:
                        new_articles.append(refreshed_art)
                    else:
                        new_articles.append(a)
                else:
                    new_articles.append(a)
            articles = new_articles
    # Save history
    history_path.parent.mkdir(parents=True, exist_ok=True)
    # Keep last 5000 records
    if len(history.get("fetched_urls", [])) > 5000:
        history["fetched_urls"] = history["fetched_urls"][-5000:]
    try:
        _safe_write_json(history_path, history, ensure_ascii=False, indent=2)
    except Exception:
        # 回退到简单写入（兼容没有 portalocker 的环境）
        with open(history_path, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)

    # 如果有已抓取的 URL，从缓存中合并条目
    already_fetched_urls = [u for u in urls if u not in new_urls]
    if already_fetched_urls:
        # 尝试从缓存匹配
        cached_matched = [a for a in cached_articles if a.get("url") in already_fetched_urls] if cached_articles else []
        if cached_matched:
            # 检查缓存条目是否需要重抓
            urls_to_refetch = [a["url"] for a in cached_matched if _needs_refetch(a)]
            if urls_to_refetch:
                print(f"  [Refetch] {len(urls_to_refetch)} cached articles missing content, attempting refetch")
                refreshed = _attempt_refetch(urls_to_refetch)
                if refreshed:
                    by_url = {a['url']: a for a in refreshed}
                    cached_matched = [
                        by_url.get(a['url'], a) for a in cached_matched
                    ]
            print(f"  [Cache] Merged {len(cached_matched)} previously fetched articles")
            articles = cached_matched + articles
        else:
            # 缓存中没有匹配项 —— history 记录了 hash 但缓存文件没有对应条目
            # 强制重新抓取这些 URL
            print(f"  [Warning] {len(already_fetched_urls)} URL(s) marked fetched but not in cache, re-fetching")
            for url in already_fetched_urls:
                try:
                    article = process_single(url, download_images=False)
                    if article and not article.get("error"):
                        standard_article = {
                            "title": article.get("title") or article.get("original_title") or "无标题",
                            "url": url,
                            "account_name": article.get("account_name", "") or article.get("author", "") or "未知公众号",
                            "content_text": article.get("content_text", "") or article.get("text", "") or article.get("full_text", ""),
                            "images": article.get("images", []),
                            "source": article.get("source", "unknown"),
                            "fetched_at": datetime.now().isoformat(),
                        }
                        articles.append(standard_article)
                        print(f"  [OK] {standard_article['title'][:40]}...")
                    time.sleep(1)
                except Exception as e:
                    print(f"  [Failed] {url[:50]}...: {e}")

    # 按输入 URL 顺序排序
    url_order = {u: i for i, u in enumerate(urls)}
    articles.sort(key=lambda a: url_order.get(a.get("url", ""), 999))

    print(f"  Fetched: {len(articles)} articles (from {len(urls)} URLs)")
    if not articles and cached_articles:
        print(f"  [Cache] Fallback to {len(cached_articles)} cached WeChat articles")
        return cached_articles
    return articles


def step2_fetch_bank_news(days: int = 7) -> list[dict]:
    """
    Step 2: Fetch bank announcements

    Args:
        days: Number of days to scrape

    Returns:
        List of announcement dicts
    """
    print("\n" + "=" * 60)
    print(f"Step 2: Fetch bank announcements (last {days} days)")
    print("=" * 60)

    try:
        # 导入 news-analyzer
        sys.path.insert(0, str(PROJECT_ROOT / "news-analyzer" / "scripts"))
        from website_scraper import scrape_bank, BANK_CONFIGS

        # 逐银行抓取，收集并返回可序列化的 dict
        all_items = []
        for bank_name, bank_cfg in BANK_CONFIGS.items():
            try:
                items = scrape_bank(bank_cfg, days=days)
                # items 可能是 CreditCardItem 对象列表，转换为 dict
                for it in items:
                    try:
                        if hasattr(it, "to_dict"):
                            all_items.append(it.to_dict())
                        elif isinstance(it, dict):
                            all_items.append(it)
                        else:
                            # 兜底转换 — 保留所有已知字段
                            all_items.append({
                                "title": getattr(it, "title", ""),
                                "bank": getattr(it, "bank", ""),
                                "category": getattr(it, "category", ""),
                                "url": getattr(it, "url", ""),
                                "raw_text": getattr(it, "raw_text", ""),
                                "images": getattr(it, "images", []),
                                "structured": getattr(it, "structured", {}),
                                "author": getattr(it, "author", ""),
                                "publish_time": getattr(it, "publish_time", ""),
                            })
                    except Exception:
                        # 忽略单条转换错误
                        continue
            except Exception as e:
                print(f"  [Warning] {bank_name} failed: {e}")

        print(f"  Total announcements: {len(all_items)}")
        return all_items

    except ImportError as e:
        print(f"  [Error] news-analyzer import failed: {e}")
        return []
    except Exception as e:
        print(f"  [Error] Scrape failed: {e}")
        return []


def step3_merge(wechat_articles: list, bank_announcements: list, batch_label: str = "") -> CreditCardBatch:
    """
    Step 3: Merge data into standard format

    使用 common/normalizer 统一标准化入口，取代手动 CreditCardItem 构造。

    Args:
        wechat_articles: WeChat article list
        bank_announcements: Bank announcement list
        batch_label: Batch label

    Returns:
        CreditCardBatch object
    """
    print("\n" + "=" * 60)
    print("Step 3: Merge data")
    print("=" * 60)

    batch = CreditCardBatch(batch_label=batch_label or get_week_label())

    # Convert WeChat articles via normalizer (支持多主题拆分)
    for article in wechat_articles:
        content_blocks = article.get("content_blocks") or []
        if content_blocks:
            # 有 content_blocks → 走多主题拆分链路
            envelope = build_article_envelope(
                url=article.get("url", ""),
                publisher_name=article.get("account_name", ""),
                publish_time=article.get("fetched_at", ""),
                raw_title=article.get("title", ""),
                content_blocks=content_blocks,
                images=article.get("images", []),
            )
            detection = detect_multi_topic(envelope)
            if detection["is_multi_topic_candidate"]:
                # 多主题 → 拆分 + 逐 topic 标准化
                topics = split_article_into_topics(envelope)
                for topic in topics:
                    item = normalize_topic(topic)
                    batch.add(item)
            else:
                # 单主题 → 走现有标准化
                item = normalize_item(article, source="wechat")
                batch.add(item)
        else:
            # 无 content_blocks → 尝试 raw_text 主题拆分兜底
            raw_text = article.get("content_text") or article.get("full_text") or article.get("text") or ""
            detection = detect_multi_topic_from_raw_text(raw_text)
            if detection["is_multi_topic"]:
                print(f"  [RawTextSplit] 文章 '{article.get('title','')[:30]}' 拆分为 {len(detection['topics'])} 个主题")
                split_confidence = detection.get("confidence", 0.75)
                split_signals = detection.get("signals", ["raw_text_bank_action_boundaries"])
                for topic in detection["topics"]:
                    topic_item_raw = _build_item_from_raw_text_topic(topic, article)
                    item = normalize_item(topic_item_raw, source="wechat")
                    # 回填多主题拆分追溯字段
                    item.is_multi_topic_split = True
                    item.topic_split_confidence = split_confidence
                    item.topic_split_signals = split_signals
                    # 低置信度标记审核
                    if split_confidence < 0.6 and "needs_topic_split_review" not in item.review_flags:
                        item.review_flags = (item.review_flags or []) + ["needs_topic_split_review"]
                    batch.add(item)
            else:
                # 单主题 → 走现有标准化
                item = normalize_item(article, source="wechat")
                batch.add(item)

    # Convert bank announcements via normalizer (keep website's category)
    for announcement in bank_announcements:
        raw = _item_to_mapping(announcement)
        item = normalize_item(
            raw, source="website",
            bank=raw.get("bank", ""),
            skip_auto_classify=True,
        )
        batch.add(item)

    print(f"  Merged: {batch.size()} items")
    print(f"  Category breakdown:")
    for cat in ["新卡", "权益变更", "活动", "公告", "其他"]:
        count = len(batch.by_category(cat))
        if count > 0:
            print(f"    {cat}: {count}")

    return batch


def step4_generate_report(batch: CreditCardBatch, output_dir: str = None) -> str | None:
    """
    Step 4: Generate Word report

    Args:
        batch: CreditCardBatch object
        output_dir: Output directory

    Returns:
        Generated docx file path, or None on failure
    """
    print("\n" + "=" * 60)
    print("Step 4: Generate Word report")
    print("=" * 60)

    output_dir = Path(output_dir or PROJECT_ROOT / "data")
    output_dir.mkdir(parents=True, exist_ok=True)

    # Save batch data
    batch_file = output_dir / "batch_merged.json"
    batch.save_json(str(batch_file))
    print(f"  Data saved: {batch_file}")

    # Generate Word report
    try:
        sys.path.insert(0, str(PROJECT_ROOT / "word-merger" / "scripts"))
        from docx_utils import safe_generate_report

        report_file = output_dir / f"Weekly_Report_{batch.batch_label}.docx"
        result = safe_generate_report(str(batch_file), str(report_file))

        if result.get('success'):
            print(f"  Report generated: {report_file}")
            return str(report_file)
        else:
            print(f"  [Error] Generation failed: {result.get('error', 'Unknown error')}")
            return None

    except Exception as e:
        print(f"  [Error] Report generation failed: {e}")
        return None


def step4_generate_md_report(batch: CreditCardBatch, output_dir: str = None) -> str | None:
    """
    Step 4b: Generate Markdown report

    Args:
        batch: CreditCardBatch object
        output_dir: Output directory

    Returns:
        Generated md file path, or None on failure
    """
    print("\n" + "=" * 60)
    print("Step 4b: Generate Markdown report")
    print("=" * 60)

    try:
        output_dir = Path(output_dir or PROJECT_ROOT / "data")
        output_dir.mkdir(parents=True, exist_ok=True)

        md_file = output_dir / f"Weekly_Report_{batch.batch_label}.md"

        # Group items by category
        categories: dict[str, list[CreditCardItem]] = {}
        for item in batch.items:
            cat = item.category or "其他"
            categories.setdefault(cat, []).append(item)

        # Category ordering
        cat_order = ["新卡", "权益变更", "活动", "公告", "其他"]
        cat_emoji = {"新卡": "🆕", "权益变更": "🔄", "活动": "🏷️", "公告": "📋", "其他": "📌"}

        lines: list[str] = []
        lines.append(f"# 信用卡周报 - {batch.batch_label}")
        lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("")

        all_images: list[tuple[str, str]] = []  # (category, alt_text, relative_path)

        for cat in cat_order:
            items = categories.get(cat, [])
            if not items:
                continue
            emoji = cat_emoji.get(cat, "📌")
            lines.append(f"## {emoji} {cat}")
            lines.append("")

            for item in items:
                raw = _item_to_mapping(item)
                title = raw.get("title", "无标题")
                highlight = raw.get("highlight_summary") or raw.get("structured", {}).get("summary", "")
                bank = raw.get("bank", "")
                source = raw.get("source", "")
                url = raw.get("url", "")
                structured = raw.get("structured", {}) or {}
                raw_text = raw.get("raw_text", "")
                images = raw.get("images", []) or []

                lines.append(f"### {title}")
                lines.append("")

                # 亮点摘要（加粗）
                if highlight:
                    lines.append(f"**亮点：** {highlight}")
                    lines.append("")

                # 来源/银行
                source_info = []
                if bank:
                    source_info.append(f"银行：{bank}")
                if source:
                    source_info.append(f"来源：{source}")
                if source_info:
                    lines.append(" | ".join(source_info))
                    lines.append("")

                # 原文链接
                if url:
                    lines.append(f"[原文链接]({url})")
                    lines.append("")

                # 结构化字段（key-value列表）
                if structured:
                    lines.append("**结构化信息：**")
                    lines.append("")
                    for key, value in structured.items():
                        if isinstance(value, (list, dict)):
                            value_str = json.dumps(value, ensure_ascii=False)
                        else:
                            value_str = str(value)
                        lines.append(f"- **{key}**：{value_str}")
                    lines.append("")

                # raw_text（截取前500字）
                if raw_text:
                    snippet = raw_text[:500]
                    if len(raw_text) > 500:
                        snippet += "..."
                    lines.append("**原文摘要：**")
                    lines.append("")
                    lines.append(snippet)
                    lines.append("")

                # images（Markdown 图片语法，相对路径）
                if images:
                    lines.append("**相关图片：**")
                    lines.append("")
                    for img in images:
                        if img.startswith("http://") or img.startswith("https://"):
                            rel_path = img
                        else:
                            try:
                                rel_path = os.path.relpath(img, str(output_dir))
                            except Exception:
                                rel_path = img
                        alt_text = f"{title} - {os.path.basename(rel_path)}".replace("\n", " ").replace("\r", "")
                        md_path = rel_path.replace("\\", "/")
                        lines.append(f"![{alt_text}]({md_path})")
                        all_images.append((cat, alt_text, rel_path))
                    lines.append("")

                lines.append("---")
                lines.append("")

        # 图片索引
        if all_images:
            lines.append("## 🖼️ 图片索引")
            lines.append("")
            for idx, (cat, alt, path) in enumerate(all_images, 1):
                alt_clean = alt.replace("\n", " ").replace("\r", "")
                path_clean = path.replace("\\", "/")
                lines.append(f"{idx}. [{alt_clean}]({path_clean}) — 分类：{cat}")
            lines.append("")

        # Write to file
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))

        print(f"  Markdown report generated: {md_file}")
        return str(md_file)

    except Exception as e:
        print(f"  [Error] Markdown report generation failed: {e}")
        return None


def step5_analyze(batch: CreditCardBatch, scorer: str = "keyword", model: str = None) -> dict:
    """
    Step 5: Card holding analysis

    Args:
        batch: CreditCardBatch object
        scorer: Scoring mode ("keyword" | "llm")
        model: LLM model name

    Returns:
        Analysis result dict
    """
    print("\n" + "=" * 60)
    print("Step 5: Card holding analysis")
    print("=" * 60)

    temp_batch = PROJECT_ROOT / "data" / "_temp_batch.json"
    temp_analysis = PROJECT_ROOT / "data" / "_temp_analysis.json"

    try:
        batch.save_json(str(temp_batch))

        script_path = PROJECT_ROOT / "card-holding-suggestion" / "scripts" / "analyze_batch.py"
        # LLM timeout: 60s 基础 + 30s per item；关键词模式默认120s
        if scorer == "llm":
            try:
                timeout_seconds = 60 + 30 * batch.size()
            except Exception:
                timeout_seconds = 60 + 30 * 10
        else:
            timeout_seconds = 120
        env = os.environ.copy()
        if model:
            env["LLM_MODEL"] = model

        command = [
            sys.executable,
            str(script_path),
            "--input",
            str(temp_batch),
            "--output",
            str(temp_analysis),
            "--scorer",
            scorer,
        ]

        print(f"  Running analyzer ({scorer}) with timeout={timeout_seconds}s")
        completed = subprocess.run(
            command,
            cwd=str(PROJECT_ROOT),
            env=env,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout_seconds,
        )

        if completed.returncode != 0:
            stderr_text = (completed.stderr or completed.stdout or "").strip()
            raise RuntimeError(f"STEP5_SUBPROCESS_FAIL: {stderr_text or f'Analyzer exited with code {completed.returncode}'}")

        if not temp_analysis.exists():
            raise RuntimeError("Analysis output file not generated")

        with open(str(temp_analysis), 'r', encoding='utf-8') as f:
            result = json.load(f)

        if completed.stderr:
            print(completed.stderr.strip())

        result.setdefault("status", "ok")
        result.setdefault("error_code", "")
        result.setdefault("error_message", "")

        # ── RAG 增强：对重点条目查询历史参考 ──
        try:
            rag_count = 0
            for cat, data in result.get('category_summary', {}).items():
                for item in data.get('items', []):
                    ev = item.get('evaluation', {})
                    if ev.get('is_highlight'):
                        refs = query_for_suggestions(
                            category=cat,
                            bank=item.get('bank', ''),
                            title=item.get('title', ''),
                            top_k=2,
                        )
                        if refs:
                            item['rag_references'] = [
                                f"{r['title']}（{r.get('date','')}）"
                                for r in refs
                            ]
                            rag_count += 1
            print(f"  RAG enriched: {rag_count} items with historical references")
        except Exception as rag_err:
            print(f"  [Warn] RAG enrichment failed: {rag_err}")

        print(f"  Analysis complete: {result.get('total_items', 0)} items")
        return result

    except subprocess.TimeoutExpired:
        print(f"  [Warn] Analysis timed out after {timeout_seconds}s")
        if scorer == "llm":
            print("  [Fallback] Retry with keyword scorer")
            try:
                fallback_result = step5_analyze(batch, scorer="keyword", model=None)
                if isinstance(fallback_result, dict):
                    fallback_result["status"] = "ok_fallback"
                    fallback_result["error_code"] = "STEP5_TIMEOUT_LLM"
                    fallback_result["error_message"] = f"LLM scorer timed out after {timeout_seconds}s; fallback to keyword succeeded"
                    fallback_result["timeout_seconds"] = timeout_seconds
                    fallback_result["fallback_scorer"] = "keyword"
                return fallback_result
            except Exception:
                return _analysis_error(
                    "STEP5_TIMEOUT_LLM",
                    f"LLM scorer timed out after {timeout_seconds}s and keyword fallback failed",
                    timeout_seconds=timeout_seconds,
                    scorer=scorer,
                    model=model or "",
                )
        return _analysis_error(
            "STEP5_TIMEOUT_KEYWORD",
            f"Keyword scorer timed out after {timeout_seconds}s",
            timeout_seconds=timeout_seconds,
            scorer=scorer,
        )

    except Exception as e:
        error_text = str(e)
        error_code = "STEP5_ANALYZE_FAILED"
        if error_text.startswith("STEP5_SUBPROCESS_FAIL:"):
            error_code = "STEP5_SUBPROCESS_FAIL"
        print(f"  [Error] Analysis failed: {e}")
        return _analysis_error(
            error_code,
            error_text,
            scorer=scorer,
            model=model or "",
        )
    finally:
        temp_batch.unlink(missing_ok=True)
        temp_analysis.unlink(missing_ok=True)


def step6_append_suggestions(report_file: str, analysis: dict) -> str:
    """
    Step 6: Save analysis results and append suggestions to Word report

    Args:
        report_file: Word report file path
        analysis: Analysis result dict from build_context()

    Returns:
        Updated docx file path
    """
    print("\n" + "=" * 60)
    print("Step 6: 追加持卡建议")
    print("=" * 60)

    if not report_file or not Path(report_file).exists():
        print("  [跳过] 报告文件不存在")
        return report_file

    if not analysis:
        print("  [跳过] 无分析结果")
        return report_file

    try:
        # 保存分析结果为 JSON
        analysis_json = Path(report_file).parent / "analysis.json"
        with open(analysis_json, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, ensure_ascii=False, indent=2)
        print(f"  分析结果已保存: {analysis_json}")

        # 追加建议到 Word 文档
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document(report_file)

            # 添加持卡建议章节
            doc.add_page_break()
            doc.add_heading('💡 持卡建议', level=1)

            error_code = analysis.get('error_code')
            if error_code:
                note = doc.add_paragraph()
                note.add_run(f'分析状态: {analysis.get("status", "unknown")}').bold = True
                note.add_run(f' | 错误代码: {error_code}')
                if analysis.get('error_message'):
                    doc.add_paragraph(f'错误信息: {analysis["error_message"]}')

            # 分类建议
            cat_summary = analysis.get('category_summary', {})
            highlighted_banks = set()
            cat_order = ['新卡', '权益变更', '活动', '公告']
            cat_emoji = {'新卡': '🆕', '权益变更': '🔄', '活动': '🏷️', '公告': '📋'}
            cat_title = {'新卡': '新卡申请建议', '权益变更': '权益变更提醒', '活动': '活动参与建议', '公告': '重要公告'}

            # 辅助函数：添加加粗字段行
            def _add_field(paragraph_title: str, value: str):
                if value:
                    p = doc.add_paragraph()
                    r = p.add_run(f'{paragraph_title}：')
                    r.bold = True
                    p.add_run(value)

            for cat in cat_order:
                if cat not in cat_summary:
                    continue
                data = cat_summary[cat]
                items = data.get('items', [])
                highlight_items = [it for it in items if it.get('evaluation', {}).get('is_highlight')]
                if not highlight_items:
                    continue

                emoji = cat_emoji.get(cat, '📌')
                doc.add_heading(f'{emoji} {cat_title.get(cat, cat)}', level=2)

                for item in highlight_items:
                    ev = item.get('evaluation', {})
                    bank = item.get('bank', '')
                    title = item.get('title', '')
                    if bank:
                        highlighted_banks.add(bank)

                    # 条目子标题：银行 · 卡名
                    sub = doc.add_paragraph()
                    r = sub.add_run(f'▸ {bank} · {title}')
                    r.bold = True
                    r.font.size = Pt(11)

                    if cat == '新卡':
                        _add_field('是否值得', ev.get('highlight_reason', '建议观望'))
                        _add_field('核心理由', ev.get('recommendation', ev.get('summary', '')))
                        _add_field('注意事项', ev.get('notes', ''))

                    elif cat == '权益变更':
                        _add_field('影响方向', ev.get('highlight_reason', '无影响'))
                        _add_field('关键变化', ev.get('recommendation', ''))
                        _add_field('应对建议', ev.get('notes', ''))

                    elif cat == '活动':
                        _add_field('是否可参与', ev.get('highlight_reason', '建议放弃'))
                        _add_field('活动价值', ev.get('activity_value', '中'))
                        _add_field('参与方式', ev.get('recommendation', ''))
                        _add_field('注意事项', ev.get('notes', ''))

                    # RAG 历史参考
                    rag_refs = item.get('rag_references', [])
                    if rag_refs:
                        ref_text = '；'.join(rag_refs[:3])
                        p = doc.add_paragraph()
                        r = p.add_run('📌 历史参考：')
                        r.bold = True
                        p.add_run(ref_text)

            # ── 综合持卡策略 ──
            doc.add_heading('📊 综合持卡策略', level=2)

            # 统计各分类高亮条目
            all_recommended = []   # 推荐/高价值
            all_watch = []         # 需关注（权益变更缩水）
            all_activity = []      # 高价值活动

            for cat in cat_order:
                if cat not in cat_summary:
                    continue
                for item in cat_summary[cat].get('items', []):
                    ev = item.get('evaluation', {})
                    if not ev.get('is_highlight'):
                        continue
                    bank = item.get('bank', '')
                    title = item.get('title', '')
                    hl_reason = ev.get('highlight_reason', '')
                    score = ev.get('overall_score', 0)
                    label = f'{bank} · {title}'

                    if cat == '新卡' and score >= 7:
                        all_recommended.append(label)
                    elif cat == '权益变更' and score <= 4:
                        all_watch.append(label)
                    elif cat == '活动' and score >= 7:
                        all_activity.append(label)

            if all_recommended:
                p = doc.add_paragraph()
                r = p.add_run('✅ 推荐申请：')
                r.bold = True
                p.add_run('、'.join(all_recommended))

            if all_activity:
                p = doc.add_paragraph()
                r = p.add_run('🎯 优先参与活动：')
                r.bold = True
                p.add_run('、'.join(all_activity))

            if all_watch:
                p = doc.add_paragraph()
                r = p.add_run('⚠️ 需要关注的权益变动：')
                r.bold = True
                p.add_run('、'.join(all_watch))

            if not (all_recommended or all_activity or all_watch):
                doc.add_paragraph('本期无特别重点推荐。建议根据个人消费习惯合理安排用卡。')

            # 涉及银行
            if highlighted_banks:
                doc.add_heading('🏦 涉及银行', level=2)
                doc.add_paragraph('、'.join(sorted(highlighted_banks)))

            doc.save(report_file)
            print(f"  持卡建议已追加到报告")

        except ImportError:
            print("  [警告] python-docx 不可用，跳过 Word 追加")
        except Exception as e:
            print(f"  [警告] 追加到 Word 失败: {e}")

        return report_file

    except Exception as e:
        print(f"  [错误] 操作失败: {e}")
        return report_file


def step7_archive(batch: CreditCardBatch, report_file: str, analysis: dict) -> str:
    """
    Step 7: Archive to knowledge base

    Args:
        batch: CreditCardBatch object
        report_file: Word report file path
        analysis: Analysis result dict

    Returns:
        Archive path
    """
    print("\n" + "=" * 60)
    print("Step 7: Archive to knowledge base")
    print("=" * 60)

    try:
        from common.archive import archive_batch

        # Find analysis JSON file
        analysis_json_path = None
        if report_file:
            candidate = Path(report_file).parent / "analysis.json"
            if candidate.exists():
                analysis_json_path = str(candidate)

        result = archive_batch(
            batch=batch,
            docx_path=report_file,
            analysis_json_path=analysis_json_path,
        )
        print(f"  Archive complete: {result}")
        return result

    except ImportError as e:
        print(f"  [Error] archive import failed: {e}")
        return ""
    except Exception as e:
        print(f"  [Error] Archive failed: {e}")
        return ""


def run_pipeline(
    wechat_urls: list[str] = None,
    bank_days: int = 7,
    scorer: str = "keyword",
    model: str = None,
    skip_fetch: bool = False,
    archive_only: bool = False,
    mode: str = 'a',
):
    """
    运行全流程（统一错误处理）

    Args:
        wechat_urls: 微信文章 URL 列表
        bank_days: 银行公告抓取天数
        scorer: 评分模式
        model: LLM 模型
        skip_fetch: 跳过抓取步骤
        archive_only: 仅归档
        mode: 运行模式 ('a'/'b' 走完整流程, 'c' 仅 Step4 生成 md 报告 → Step7 归档)
    """
    from common.errors import PipelineResult

    result = PipelineResult("信用卡周报管道")

    print("=" * 60)
    print("Credit Card Weekly Report Pipeline")
    print(f"Execution time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 60)

    batch_label = get_week_label()

    if archive_only:
        batch_file = PROJECT_ROOT / "data" / "batch_merged.json"
        if batch_file.exists():
            batch = CreditCardBatch.load_json(str(batch_file))
            step7_archive(batch, "", {})
        else:
            print("[Error] batch_merged.json not found")
        return

    # Step 1-2: 抓取数据（每一步失败均不影响后续，默认值为空列表）
    wechat_articles: list[dict] = []
    bank_announcements: list[dict] = []

    if not skip_fetch:
        wechat_articles = result.run_step(
            "Step1 微信抓取",
            lambda: step1_fetch_wechat(wechat_urls) or [],
            default=[],
        )
        bank_announcements = result.run_step(
            "Step2 银行公告",
            lambda: step2_fetch_bank_news(bank_days) or [],
            default=[],
        )
    else:
        print("\n[Skip] Fetch steps")

    # Step 3: 合并
    batch = result.run_step(
        "Step3 合并数据",
        lambda: step3_merge(wechat_articles, bank_announcements, batch_label),
        default=CreditCardBatch(batch_label=batch_label),
    )

    if batch.size() == 0:
        print("\n[End] No data to process")
        result.print_summary()
        return

    # Step 3.5: LLM 审核（对低置信度条目做语义修正）
    try:
        from common.llm_review import llm_review_items, apply_suggestions
        from common.review import build_review_queue

        review_queue = build_review_queue(batch.items)
        flagged = review_queue.get("flagged_items", [])
        high_medium = [f for f in flagged if f.get("severity") in ("high", "medium")]

        if high_medium:
            print(f"\n  [LLM Review] {len(high_medium)} items need review")
            suggestions = llm_review_items(high_medium)
            if suggestions:
                modified = apply_suggestions(batch.items, suggestions)
                if modified:
                    print(f"  [LLM Review] Applied {modified} corrections")
                    # 重新生成展示字段（分类/银行/标题可能已修正）
                    from common.normalizer import _compute_confidence
                    from common.display_fields import generate_display_fields
                    from common.review import generate_review_flags
                    for item in batch.items:
                        for s in suggestions:
                            if s["item_id"] == item.item_id and s["suggestion"]:
                                dr = generate_display_fields(
                                    bank=item.bank, category=item.category,
                                    structured=item.structured,
                                    structured_clean=item.structured_clean,
                                    raw_title=item.raw_title,
                                    raw_text=item.raw_text,
                                )
                                item.title = dr["title"]
                                item.display_title = dr["display_title"]
                                item.highlight_summary = dr["highlight_summary"]
                                item.confidence = _compute_confidence(
                                    item.category_candidates, 0.0,
                                    dr["title_source"], item.raw_title,
                                    item.structured, item.category,
                                )
                                item.review_flags = generate_review_flags(item)
                                break
        else:
            print("\n  [LLM Review] No high/medium items, skipping")
    except Exception as e:
        print(f"\n  [LLM Review] Skipped: {e}")

    # Step 4: 生成报告
    if mode == 'c':
        report_file = result.run_step(
            "Step4 生成 MD 报告",
            lambda: step4_generate_md_report(batch),
            default=None,
        )
    else:
        report_file = result.run_step(
            "Step4 生成报告",
            lambda: step4_generate_report(batch),
            default=None,
        )

    # Step 5: 分析（mode c 跳过）
    if mode == 'c':
        analysis = {}
        print("\n[Skip] Step5 (mode c)")
    else:
        analysis = result.run_step(
            "Step5 持卡分析",
            lambda: step5_analyze(batch, scorer, model),
            default={},
        )

    # Step 6: 追加持卡建议（mode c 跳过，需要 report_file 存在）
    if mode == 'c':
        print("\n[Skip] Step6 (mode c)")
    elif report_file:
        result.run_step(
            "Step6 追加建议",
            lambda: step6_append_suggestions(report_file, analysis),
            default=report_file,
        )
    else:
        print("\n[Skip] Step6 (no report file)")

    # Step 7: 归档
    result.run_step(
        "Step7 归档",
        lambda: step7_archive(batch, report_file or "", analysis or {}),
    )

    result.print_summary()

    # 导出审核队列
    if batch and batch.size() > 0:
        try:
            review_dir = str(PROJECT_ROOT / "data" / "review")
            json_path, md_path = export_review_queue(
                batch.items,
                review_dir,
                batch_label=batch.batch_label,
            )
            flagged = sum(1 for it in batch.items if it.review_flags)
            print(f"  Review queue: {flagged}/{batch.size()} flagged → data/review/{os.path.basename(json_path)}")
        except Exception as e:
            print(f"  [Warn] Review queue export failed: {e}")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="信用卡周报全流程 Agent",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 完整流程
  python _agent.py

  # 指定微信文章
  python _agent.py --wechat-url https://mp.weixin.qq.com/s/xxx

  # 跳过抓取，使用已有数据
  python _agent.py --skip-fetch

  # 仅归档
  python _agent.py --archive-only

  # 使用 LLM 评分
  python _agent.py --scorer llm --model deepseek-chat
        """
    )

    parser.add_argument("--wechat-url", nargs="+", help="微信文章 URL 列表")
    parser.add_argument("--bank-days", type=int, default=7, help="银行公告抓取天数 (默认 7)")
    parser.add_argument("--scorer", choices=["keyword", "llm"], default="keyword",
                        help="评分模式 (默认 keyword)")
    parser.add_argument("--model", help="LLM 模型名称")
    parser.add_argument("--skip-fetch", action="store_true", help="跳过抓取步骤")
    parser.add_argument("--archive-only", action="store_true", help="仅归档")
    parser.add_argument("--mode", choices=["a", "c"], default="a",
                        help="运行模式: a 完整流程, c 仅生成 md 报告 + 归档 (默认 a)")

    args = parser.parse_args()

    run_pipeline(
        wechat_urls=args.wechat_url,
        bank_days=args.bank_days,
        scorer=args.scorer,
        model=args.model,
        skip_fetch=args.skip_fetch,
        archive_only=args.archive_only,
        mode=args.mode,
    )


if __name__ == "__main__":
    main()
