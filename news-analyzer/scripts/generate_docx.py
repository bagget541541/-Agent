#!/usr/bin/env python3
"""
Word文档生成脚本
功能：将公告分析结果生成为格式化的Word文档，支持图片嵌入与LLM点评。
"""

import json
import sys
import os
import argparse
from datetime import datetime
from typing import List, Dict


def setup_docx():
    """导入python-docx模块，返回常用类和工具函数。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return (Document, Pt, RGBColor, Inches, WD_ALIGN_PARAGRAPH, qn, OxmlElement)
    except ImportError:
        print(json.dumps({"error": "python-docx未安装，请运行: pip install python-docx"}, ensure_ascii=False))
        sys.exit(1)


def _download_image(img_url: str, save_dir: str) -> str | None:
    """从 URL 下载图片到本地，返回本地路径。"""
    try:
        import hashlib
        from urllib.parse import urlparse
        import requests
        os.makedirs(save_dir, exist_ok=True)
        resp = requests.get(img_url, timeout=30, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        resp.raise_for_status()
        ext = os.path.splitext(urlparse(img_url).path)[1] or '.jpg'
        name = hashlib.md5(img_url.encode()).hexdigest()[:12] + ext
        path = os.path.join(save_dir, name)
        with open(path, 'wb') as f:
            f.write(resp.content)
        return path
    except Exception as e:
        print(f"  [图片下载失败] {img_url}: {e}", file=sys.stderr)
        return None


def _set_east_asia(run, font_name: str, qn, OxmlElement):
    """为 run 同步设置东亚字体。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _embed_image(doc, img_path: str, qn, OxmlElement):
    """在文档中嵌入一张图片（支持本地路径与下载后的 URL 图片）。"""
    if not os.path.isfile(img_path):
        print(f"  [图片不存在] {img_path}", file=sys.stderr)
        return
    try:
        from docx.shared import Cm, Inches
        from PIL import Image as PILImage
        from io import BytesIO

        pil_img = PILImage.open(img_path)
        w, h = pil_img.size

        # 计算合适的渲染尺寸
        max_w_inch = 4.5
        max_h_inch = 4.0
        dpi = 96
        w_inch = w / dpi
        h_inch = h / dpi

        if w_inch > max_w_inch:
            ratio = max_w_inch / w_inch
            w_inch = max_w_inch
            h_inch *= ratio
        if h_inch > max_h_inch:
            ratio = max_h_inch / h_inch
            h_inch = max_h_inch
            w_inch *= ratio

        para = doc.add_paragraph()
        para.alignment = 1  # CENTER
        run = para.add_run()
        run.add_picture(img_path, width=Inches(w_inch), height=Inches(h_inch))
    except Exception as e:
        print(f"  [图片嵌入失败] {img_path}: {e}", file=sys.stderr)


def generate_word_document(data: List[Dict], output_path: str) -> bool:
    """
    生成Word文档

    Args:
        data: 公告数据列表，每个元素可包含 title, url, message, comment, images
        output_path: 输出文件路径

    Returns:
        是否成功生成文档
    """
    Document, Pt, RGBColor, Inches, WD_ALIGN_PARAGRAPH, qn, OxmlElement = setup_docx()
    doc = Document()

    # 设置默认字体 + 东亚字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    _rPr = style.element.get_or_add_rPr()
    _rFonts = _rPr.find(qn('w:rFonts'))
    if _rFonts is None:
        _rFonts = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:rFonts')
        _rPr.insert(0, _rFonts)
    _rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 文档标题
    title = doc.add_heading('公告分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        _set_east_asia(run, 'Microsoft YaHei', qn, OxmlElement)

    # 生成时间
    time_paragraph = doc.add_paragraph()
    time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_run = time_paragraph.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    time_run.font.size = Pt(10)
    time_run.font.color.rgb = RGBColor(128, 128, 128)
    _set_east_asia(time_run, 'Microsoft YaHei', qn, OxmlElement)

    # 添加分隔线
    doc.add_paragraph('_' * 80)

    # 准备图片临时目录（与输出文档同目录）
    img_dir = os.path.join(os.path.dirname(output_path) or '.', '.images_temp')
    os.makedirs(img_dir, exist_ok=True)

    # 为每个公告添加内容
    for idx, item in enumerate(data, 1):
        print(f"正在处理第 {idx} 条公告...", file=sys.stderr)
        print(f"标题: {item.get('title', '无')}", file=sys.stderr)
        msg_preview = (item.get('message', '') or '')[:50]
        print(f"消息内容: {msg_preview if msg_preview else '空'}...", file=sys.stderr)

        # 公告标题
        heading = doc.add_heading(f'{idx}. {item.get("title", "无标题")}', level=2)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            _set_east_asia(run, 'Microsoft YaHei', qn, OxmlElement)

        # 原文链接
        link_paragraph = doc.add_paragraph()
        link_run = link_paragraph.add_run(f'原文链接：{item.get("url", "")}')
        link_run.font.size = Pt(9)
        link_run.font.color.rgb = RGBColor(0, 102, 204)
        _set_east_asia(link_run, 'Microsoft YaHei', qn, OxmlElement)

        # 消息内容
        doc.add_heading('消息内容', level=3)
        message_text = item.get("message", "")
        if not message_text:
            message_text = "（无消息内容）"
        message_paragraph = doc.add_paragraph(message_text)
        message_paragraph.paragraph_format.left_indent = Inches(0.25)
        for run in message_paragraph.runs:
            _set_east_asia(run, 'Microsoft YaHei', qn, OxmlElement)

        # ── 图片嵌入（新增） ──
        images = item.get("images", [])
        if images:
            print(f"  图片 × {len(images)}", file=sys.stderr)
            for img in images:
                if not img:
                    continue
                img = str(img).strip()
                local_path = None
                if img.startswith(('http://', 'https://')):
                    local_path = _download_image(img, img_dir)
                else:
                    local_path = img if os.path.isfile(img) else None
                if local_path:
                    _embed_image(doc, local_path, qn, OxmlElement)

        # ── LLM 点评（新增） ──
        comment = item.get("comment", "")
        if comment:
            doc.add_heading('AI 点评', level=3)
            comment_para = doc.add_paragraph(comment)
            comment_para.paragraph_format.left_indent = Inches(0.25)
            # 点评用灰色蓝字体
            for run in comment_para.runs:
                run.font.color.rgb = RGBColor(51, 102, 153)
                _set_east_asia(run, 'Microsoft YaHei', qn, OxmlElement)

        # 公告之间添加空行（最后一个除外）
        if idx < len(data):
            doc.add_paragraph()

    # 保存文档
    doc.save(output_path)
    return True


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='生成公告分析Word文档')
    parser.add_argument('--input', type=str, required=True,
                        help='输入JSON文件路径，包含公告分析数据')
    parser.add_argument('--output', type=str, default='公告分析报告.docx',
                        help='输出Word文件路径（默认: 公告分析报告.docx）')

    args = parser.parse_args()

    # 读取输入数据
    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(json.dumps({"error": f"找不到输入文件: {args.input}"}, ensure_ascii=False))
        sys.exit(1)
    except json.JSONDecodeError:
        print(json.dumps({"error": "输入文件格式错误，应为有效的JSON"}, ensure_ascii=False))
        sys.exit(1)

    # 生成Word文档
    try:
        if not isinstance(data, list):
            print(json.dumps({"error": "输入数据应为JSON数组"}, ensure_ascii=False))
            sys.exit(1)
        generate_word_document(data, args.output)
        print(json.dumps({
            "success": True,
            "output": args.output,
            "count": len(data)
        }, ensure_ascii=False, indent=2))
    except Exception as e:
        print(json.dumps({"error": f"生成文档失败: {str(e)}"}, ensure_ascii=False))
        sys.exit(1)


if __name__ == '__main__':
    main()
